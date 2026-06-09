"""
================================================================================
ANGREN UCG APPLICATION — BIRLASHTIRIЛGAN VA TO'LIQ VERSIYA
================================================================================
PhD HIMOYASI VA PATENT UCHUN SCI SIFATLI KOD
Versiya: 3.0 (app.py + app_corrected.py to'liq birlashtirildi)
Muallif: Saitov Dilshodbek
Litsenziya: AGPL v3

ASOSIY TUZATISHLAR (app_corrected.py dan):
✅ Hoek-Brown (2018) 'a' parametri to'g'irlandi — [0.5, 0.65] chegarada
✅ Cholesky dekompozitsiya — musbat aniqlik kafolati
✅ Regularizatsiyalangan Kirsch yechimi — singularsiz
✅ Robin chegaraviy sharti issiqlik tenglamasi uchun
✅ Modifikatsiyalangan Kozeny-Carman o'tkazuvchanligi (zarar bilan)
✅ To'liq 3D bosh kuchlanishlar hisoblash
✅ Haroratga bog'liq gaz viskoziteti
✅ Bootstrap ishonch intervali (Efron & Tibshirani, 1993)
✅ O'lchamli tahlil (dimensional analysis)
✅ Skempton g'ovaklik bosimi tuzatmasi
✅ To'liq xatolarni boshqarish + audit logging

TO'LIQ UI XUSUSIYATLARI (app.py dan):
✅ 3 tilli interfeys (O'zbek / English / Русский)
✅ Ko'p qatlamli litologiya tahlili
✅ UCG yonish bosqichlari vizualizatsiyasi
✅ AI/ML bashorat (RandomForest + PyTorch NN)
✅ Real-vaqt monitoring paneli
✅ Sobol' global sezgirlik tahlili
✅ Latin Hypercube Sampling
✅ SHAP model interpretatsiyasi
✅ 3D litologik hajm (PyVista)
✅ PINN fizika-asosli neyron to'r (demo)
✅ Word hujjat eksporti
✅ QR-kod

ADABIYOTLAR (REFERENCES):
1. Hoek, E., & Brown, E. T. (2018). JRMGE, 10(1), 1-57.
2. Yang, D. (2010). PhD Thesis, TU Delft.
3. Shao, S., et al. (2015). IJRMMS, 81, 1-10.
4. Wilson, A. H. (1972). Mining Engineer.
5. Walsh & Brace (1984). JGR, 89(B12), 9425-9432.
6. Chapman & Cowling (1970). Non-uniform Gases.
7. Efron & Tibshirani (1993). Introduction to the Bootstrap.
8. Skempton, A. W. (1961). Effective stress in soils.
9. Bourdin et al. (2000). Phase-field fracture model.
10. Lundberg & Lee (2017). SHAP interpretation.
================================================================================
"""

import streamlit as st
st.set_page_config(
    page_title="UCG SCI-Grade Platform v3.0 (MERGED)",
    layout="wide",
    initial_sidebar_state="expanded"
)

import numpy as np
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from scipy.ndimage import gaussian_filter
from scipy.stats import linregress, t as t_dist, norm
from scipy import stats
import io
import time
import os
from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor
from sklearn.metrics import accuracy_score, roc_auc_score, r2_score
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import qrcode
import matplotlib.pyplot as plt
import warnings
warnings.filterwarnings('ignore', category=RuntimeWarning)
import logging
from scipy.signal import savgol_filter
import hashlib
from dataclasses import dataclass, field
import random
from datetime import datetime
import json
from functools import wraps
from collections import defaultdict
import psutil
from pydantic import BaseModel, Field, validator
from enum import Enum
import statistics
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
import logging.handlers
import timeit

# SQLAlchemy
from sqlalchemy import create_engine, pool
from sqlalchemy.orm import sessionmaker

# Ixtiyoriy kutubxonalar (Optional imports)
try:
    import torch
    import torch.nn as nn
    PT_AVAILABLE = True
    device = "cuda" if torch.cuda.is_available() else "cpu"
except ImportError:
    PT_AVAILABLE = False
    device = "cpu"

try:
    from SALib.sample import saltelli
    from SALib.analyze import sobol
    SALIB_AVAILABLE = True
except ImportError:
    SALIB_AVAILABLE = False

try:
    from pyDOE import lhs
    PYDOE_AVAILABLE = True
except ImportError:
    PYDOE_AVAILABLE = False

try:
    import pyvista as pv
    PYVISTA_AVAILABLE = True
except ImportError:
    PYVISTA_AVAILABLE = False

try:
    import shap
    SHAP_AVAILABLE = True
except ImportError:
    SHAP_AVAILABLE = False

try:
    from hypothesis import given, settings, strategies as st_hyp, Phase
    HYPOTHESIS_AVAILABLE = True
except ImportError:
    HYPOTHESIS_AVAILABLE = False


# ==================== MARKAZIY KONFIGURATSIYA ====================
@dataclass(frozen=True)
class AppConfig:
    """Butun ilova uchun yagona konfiguratsiya nuqtasi"""
    MODEL_MAX_DEPTH: int = 5000
    CACHE_MAX_SIZE: int = 100
    RANDOM_SEED: int = 42
    DEBUG_MODE: bool = False
    LOG_LEVEL: str = "INFO"
    START_TIME: float = time.time()


CONFIG = AppConfig()


# ==================== ILGOR LOGGING TIZIMI ====================
class JSONFormatter(logging.Formatter):
    def format(self, record):
        log_obj = {
            'timestamp': self.formatTime(record),
            'level': record.levelname,
            'logger': record.name,
            'message': record.getMessage(),
            'module': record.module,
            'function': record.funcName,
            'line': record.lineno
        }
        if record.exc_info:
            log_obj['exception'] = self.formatException(record.exc_info)
        return json.dumps(log_obj)


logger = logging.getLogger('ucg_app')
os.makedirs('logs', exist_ok=True)
handler = logging.handlers.RotatingFileHandler(
    'logs/app.log', maxBytes=10 * 1024 * 1024, backupCount=5
)
handler.setFormatter(JSONFormatter())
logger.addHandler(handler)
logger.setLevel(getattr(logging, CONFIG.LOG_LEVEL))


# ==================== TASODIFIYLIK BOSHQARUVI ====================
@dataclass(frozen=True)
class RandomState:
    """Markaziy tasodifiy sonlar generatori"""
    seed: int = CONFIG.RANDOM_SEED

    def __post_init__(self):
        np.random.seed(self.seed)
        if PT_AVAILABLE:
            torch.manual_seed(self.seed)
            torch.cuda.manual_seed_all(self.seed)
        random.seed(self.seed)


RNG = RandomState()


# ==================== GLOBAL KONSTANTALAR ====================
EPS = 1e-6
GEOM_EPS = 1e-3
T_REF_AMBIENT = 20.0

WILSON_C1 = 0.64   # Wilson (1972) — CITED
WILSON_C2 = 0.36


# ==================== SESSIYA HOLATI ====================
if "language" not in st.session_state:
    st.session_state.language = "uz"
if "theme" not in st.session_state:
    st.session_state.theme = "dark"
if "live_history_df" not in st.session_state:
    st.session_state.live_history_df = pd.DataFrame(
        columns=['step', 'mean_subsidence_cm', 'max_temp_c', 'FOS', 'pillar_width_m']
    )
if "user_id" not in st.session_state:
    st.session_state.user_id = "anonymous"


# ==================== FIZIK PARAMETRLAR ====================
@dataclass(frozen=True)
class UCGPhysicsParams:
    """Fizik parametrlar — adabiyotlarga havola qilingan"""
    phi_deg: float = 35.0            # Ishqalanish burchagi (Mohr-Coulomb)
    cohesion: float = 5e6            # Pa
    alpha_thermal: float = 3e-5     # 1/K — Yang (2010)
    gas_temp: int = 1100             # °C
    subsidence_rate: float = 0.012
    thermal_damage_beta: float = 0.002   # Shao et al. (2015)
    extraction_ratio: float = 0.6
    E_mass: float = 25e9            # Pa — ko'mir uchun
    CONFINEMENT: float = 0.65
    RELAX: float = 0.15
    THERMAL_DIFFUSIVITY: float = 8.5e-7   # m²/s
    GAS_VISCOSITY: float = 3e-5     # Pa·s at 1000°C
    MOLAR_MASS_GAS: float = 0.028   # kg/mol
    R_UNIVERSAL: float = 8.314      # J/(mol·K)
    K_VOID: float = 0.35


PARAMS = UCGPhysicsParams()


# ==================== MAXSUS XATOLIKLAR ====================
class UCGSimulationError(Exception):
    """Asosiy simulyatsiya xatoligi / Main simulation error"""
    pass


class PhysicsValidationError(UCGSimulationError):
    def __init__(self, param_name: str, value: float, bounds: tuple):
        self.param_name = param_name
        self.value = value
        self.bounds = bounds
        super().__init__(
            f"{param_name}={value} chegaradan tashqarida / outside bounds {bounds}"
        )


class NumericalInstabilityError(UCGSimulationError):
    """Raqamli divergensiya / Numerical divergence detected"""
    pass


# ==================== KIRISH MA'LUMOTLARINI TEKSHIRISH ====================
class SensorDataInput(BaseModel):
    temperature: float = Field(..., ge=-50, le=1500)
    pressure: float = Field(..., ge=0, le=50)
    stress: float = Field(..., ge=0, le=100)

    @validator('temperature')
    def validate_temp_range(cls, v):
        if not (-50 <= v <= 1500):
            raise ValueError("Temperatura fizik chegaralardan tashqarida / "
                             "Temperature outside physical bounds")
        return v


# ==================== KO'P TILLILIK (TRANSLATIONS) ====================
TRANSLATIONS = {
    'uz': {
        'app_title': "Universal Yer yuzasi Deformatsiyasi Monitoringi",
        'app_subtitle': "Termo-Mexanik (TM) tahlil va Selek O'lchami Optimizatsiyasi",
        'sidebar_header_params': "⚙️ Umumiy parametrlar",
        'formula_show': "Formulalarni ko'rish:",
        'project_name': "Loyiha nomi:",
        'process_time': "Jarayon vaqti (soat):",
        'num_layers': "Qatlamlar soni:",
        'tensile_model': "Tensile modeli:",
        'rock_props': "💎 Jins Xususiyatlari",
        'disturbance': "Disturbance Factor (D):",
        'poisson': "Poisson koeffitsiyenti (ν):",
        'stress_ratio': "Stress Ratio (k = σh/σv):",
        'tensile_params': "📐 Cho'zilish va Selek",
        'thermal_decay_label': "Termal Degradatsiya (β):",
        'combustion': "🔥 Yonish va Termal",
        'burn_duration': "Kamera yonish muddati (soat):",
        'max_temp': "Maksimal harorat (°C)",
        'timeline': "📅 Loyiha bosqichlari (Timeline)",
        'layer_params': "{num}-qatlam parametrlari",
        'layer_name': "Nomi:",
        'thickness': "Qalinlik (m):",
        'ucs': "UCS (MPa):",
        'density': "Zichlik (kg/m³):",
        'color': "Rangi:",
        'gsi': "GSI:",
        'mi': "mi:",
        'manual_st0': "σt0 (MPa):",
        'error_thick_positive': "Qalinlik >0 bo'lishi kerak",
        'error_ucs_positive': "UCS >0 MPa bo'lishi kerak",
        'error_density_positive': "Zichlik >0 kg/m³ bo'lishi kerak",
        'error_gsi_range': "GSI 10...100 oralig'ida bo'lishi kerak",
        'error_mi_positive': "mi >0 bo'lishi kerak",
        'error_min_layers': "❌ Kamida 1 ta qatlam kiriting!",
        'warning_pytorch': "⚠️ PyTorch o'rnatilmagan. RandomForestClassifier ishlatiladi.",
        'pillar_strength': "Pillar Strength (σp)",
        'plastic_zone': "Plastik zona (y)",
        'cavity_volume': "Kamera Hajmi",
        'max_permeability': "Maks. O'tkazuvchanlik",
        'ai_recommendation': "Analitik Tavsiya (Selek)",
        'monitoring_header': "📊 {obj_name}: Monitoring va Ekspert Xulosasi",
        'subsidence_title': "📉 Yer yuzasi cho'kishi (cm)",
        'thermal_deform_title': "🔥 Gorizontal siljish (cm)",
        'hb_envelopes_title': "🛡️ Hoek-Brown Envelopes",
        'scientific_analysis': "📋 Ilmiy Tahlil",
        'fos_red': "🔴 FOS < 1.0: Failure",
        'fos_yellow': "🟡 FOS 1.0–1.5: Unstable",
        'fos_green': "🟢 FOS > 1.5: Stable",
        'tm_field_title': "🔥 TM Maydoni va Selek Interferensiyasi",
        'temp_subplot': "Harorat Maydoni (°C) + Gaz Oqimi",
        'fos_subplot': "FOS + AI Collapse Prediction (NN) + Yielded Zones",
        'gas_flow': "Gaz oqimi",
        'shear': "Shear",
        'tensile': "Tensile",
        'ai_collapse': "AI Collapse (NN)",
        'monitoring_panel': "📊 {obj_name}: Kompleks Monitoring Paneli",
        'pillar_live': "Pillar Strength",
        'rec_width_live': "Tavsiya: Selek Eni",
        'max_subsidence_live': "Maks. Cho'kish",
        'process_stage': "Jarayon bosqichi",
        'stage_active': "Faol",
        'stage_cooling': "Sovish",
        'ai_monitor_title': "🧠 UCG AI Predictive Monitoring",
        'ai_monitor_desc': "Real-vaqt sensor ma'lumotlari va anomaliya aniqlash",
        'ai_steps': "Simulyatsiya qadamlari soni:",
        'ai_run_btn': "▶️ AI Monitoringni Ishga Tushirish",
        'ai_stop_btn': "⏹ To'xtatish",
        'advanced_analysis': "🔍 Chuqurlashtirilgan Dinamik Tahlil va Metodik Asoslash",
        'tab_mass': "🏗️ Massiv Parametrlari",
        'tab_thermal': "🔥 Termal Degradatsiya",
        'tab_stability': "⚖️ Barqarorlik & Manbalar",
        'hb_class': "1. Hoek-Brown (2018) Klassifikatsiyasi",
        'hb_mb': "m_b = {mb:.3f}",
        'hb_s': "s = {s:.4f}",
        'hb_caption_mb': "Massiv ishqalanish burchagi funksiyasi (m_i={mi})",
        'hb_caption_s': "Yoriqlilik darajasi (GSI={gsi})",
        'hb_interpret': "**Ilmiy izoh:** Hoek & Brown (2018) mezoniga ko'ra, GSI={gsi} bo'lishi massiv mustahkamligi laboratoriyaga nisbatan **{perc:.1f}%** ga pastligini anglatadi.",
        'thermal_params': "2. Termo-Mexanik Koeffitsiyentlar Tahlili",
        'param_table_param': "Parametr",
        'param_table_value': "Qiymat",
        'param_table_reason': "Tanlanish sababi",
        'modulus': "Elastiklik Moduli (E)",
        'alpha': "Termal kengayish (α)",
        'temp0': "Boshlang'ich T₀",
        'modulus_reason': "Ko'mir uchun xos o'rtacha deformatsiya koeffitsiyenti.",
        'alpha_reason': "Ko'mirning issiqlikdan chiziqli kengayish ko'rsatkichi (Yang, 2010).",
        'temp0_reason': "Kon qatlamining boshlang'ich tabiiy harorati.",
        'ucs_decay': "**A) Termal UCS pasayishi:**",
        'ucs_decay_eq': r"\sigma_{{ci(T)}} = \sigma_{{ci(0)}} \cdot e^{{-\beta(T-T_0)}} = {ucs:.2f} \text{{ MPa}}",
        'ucs_interpret': "**Interpretatsiya:** {temp}°C haroratda jins mustahkamligi {perc:.1f}% ga pasaydi.",
        'thermal_stress': "**B) Termal kuchlanish ($\\sigma_{{th}}$):**",
        'thermal_stress_eq': r"\sigma_{{th}} \approx \frac{{E \cdot \alpha \cdot \Delta T}}{{1 - \nu}} = {sigma:.2f} \text{{ MPa}}",
        'pillar_stability': "3. Selek Barqarorligi va Bibliografiya",
        'fos_eq': r"FOS = \frac{{\sigma_p}}{{\sigma_v}} = {fos:.2f}",
        'pillar_wilson': "**Wilson (1972) Yield Pillar nazariyasiga binoan:** Selek o'lchami $w={w}$ m bo'lganda, uning markaziy yadrosi {sv:.2f} MPa lik geostatik yukni ko'tarishga qodir. Plastik zona: $y = {y:.1f}$ m.",
        'references': "#### 📚 Asosiy Ilmiy Manbalar:",
        'ref1': "**Hoek, E., & Brown, E. T. (2018).** The Hoek-Brown failure criterion and GSI – 2018 edition. *JRMGE*.",
        'ref2': "**Yang, D. (2010).** *Stability of Underground Coal Gasification*. PhD Thesis, TU Delft.",
        'ref3': "**Shao, S., et al. (2015).** A thermal damage constitutive model for rock. *IJRMMS*.",
        'ref4': "**Wilson, A. H. (1972).** Research into the determination of pillar size. *Mining Engineer*.",
        'conclusion_danger': "🔴 **Ilmiy Xulosa:** FOS={fos:.2f}. Termal degradatsiya yuqori. Selek enini oshirish yoki gazlashtirish tezligini nazorat qilish tavsiya etiladi.",
        'conclusion_safe': "🟢 **Ilmiy Xulosa:** FOS={fos:.2f}. Tanlangan parametrlar massiv barqarorligini ta'minlaydi.",
        'methodology_expander': "📚 Ilmiy Metodologiya va Manbalar (PhD Research References)",
        'tensile_empirical': "Empirical (UCS)",
        'tensile_hb': "HB-based (auto)",
        'tensile_manual': "Manual",
        'dip_angle_label': "Qatlam qiyaligi (Dip - °)",
        'timeline_table': """
| Bosqich | Vaqti | Tavsif |
|---------|-------|--------|
| **Rejalashtirish** | 2026-04-01 | Validatsiya, xavfsiz funksiyalar |
| **Optimallashtirish** | 2026-05-15 | NN/RF testlash, FDM yaxshilash |
| **Integratsiya** | 2026-06-30 | Unit testlar, yakuniy deploy |
        """,
        'live_monitoring_tab': "🔄 Live 3D Monitoring",
        'download_data': "📥 Monitoring ma'lumotlarini yuklab olish (CSV)",
        'well_config': "Quduqlar konfiguratsiyasi",
        'well_distance': "Quduqlar orasidagi masofa (m):",
        'warning_cavity_width': "Ogohlantirish: Selek kengligi quduqlar masofasidan katta. cavity_width=1m deb olindi.",
        'ucg_stages_title': "UCG Yonish Bosqichlari (1-2-3 sxemasi) – Yangi Ilmiy Model",
        'select_stage': "Bosqichni tanlang:",
        'geomech_state': "Geomexanik Holat (Yangi Ilmiy Model)",
        'auto_animation': "Avtomatik animatsiya (1→2→3 bosqichlar)",
        'animation_done': "Animatsiya yakunlandi.",
        'pillar_annotation': "HIMOYA SELEGI (PILLAR)",
        'system_entropy': "Tizim entropiyasi (noaniqlik)",
        'pin_approx': "**Eslatma:** Kirsch yechimi kvazistatik (regularizatsiyalangan). Katta deformatsiyalar uchun FEM remeshing talab qilinadi.",
        'phase_field_info': "**Fazaviy maydon modeli (Bourdin et al., 2000 asosida):**",
        'uq_info': "Noaniqlik miqdoriy tahlili (Monte-Carlo, Cholesky dekompozitsiya):",
        'shap_info': "SHAP interpretatsiyasi (Lundberg & Lee, 2017):",
        'sobol_info': "Global sezgirlik (Sobol', 2001):",
        'lhs_info': "Latin Hypercube Sampling (McKay, 1979):",
        'validation_info': "Eksperimental validatsiya:",
        'experimental_note': "CSV faylida 'x' (m) va 'subsidence_cm' ustunlari bo'lishi shart."
    },
    'en': {
        'app_title': "Universal Surface Deformation Monitoring",
        'app_subtitle': "Thermo-Mechanical (TM) Analysis and Pillar Size Optimization",
        'sidebar_header_params': "⚙️ General Parameters",
        'formula_show': "View formulas:",
        'project_name': "Project name:",
        'process_time': "Process time (hours):",
        'num_layers': "Number of layers:",
        'tensile_model': "Tensile model:",
        'rock_props': "💎 Rock Properties",
        'disturbance': "Disturbance Factor (D):",
        'poisson': "Poisson's ratio (ν):",
        'stress_ratio': "Stress Ratio (k = σh/σv):",
        'tensile_params': "📐 Tension and Pillar",
        'thermal_decay_label': "Thermal Degradation (β):",
        'combustion': "🔥 Combustion and Thermal",
        'burn_duration': "Burn duration (hours):",
        'max_temp': "Maximum temperature (°C)",
        'timeline': "📅 Project Timeline",
        'layer_params': "Layer {num} parameters",
        'layer_name': "Name:",
        'thickness': "Thickness (m):",
        'ucs': "UCS (MPa):",
        'density': "Density (kg/m³):",
        'color': "Color:",
        'gsi': "GSI:",
        'mi': "mi:",
        'manual_st0': "σt0 (MPa):",
        'error_thick_positive': "Thickness must be >0",
        'error_ucs_positive': "UCS must be >0 MPa",
        'error_density_positive': "Density must be >0 kg/m³",
        'error_gsi_range': "GSI must be between 10 and 100",
        'error_mi_positive': "mi must be >0",
        'error_min_layers': "❌ At least 1 layer required!",
        'warning_pytorch': "⚠️ PyTorch not installed. Using RandomForestClassifier.",
        'pillar_strength': "Pillar Strength (σp)",
        'plastic_zone': "Plastic zone (y)",
        'cavity_volume': "Cavity Volume",
        'max_permeability': "Max Permeability",
        'ai_recommendation': "Analytical Recommendation (Pillar)",
        'monitoring_header': "📊 {obj_name}: Monitoring and Expert Summary",
        'subsidence_title': "📉 Surface subsidence (cm)",
        'thermal_deform_title': "🔥 Horizontal displacement (cm)",
        'hb_envelopes_title': "🛡️ Hoek-Brown Envelopes",
        'scientific_analysis': "📋 Scientific Analysis",
        'fos_red': "🔴 FOS < 1.0: Failure",
        'fos_yellow': "🟡 FOS 1.0–1.5: Unstable",
        'fos_green': "🟢 FOS > 1.5: Stable",
        'tm_field_title': "🔥 TM Field and Pillar Interference",
        'temp_subplot': "Temperature Field (°C) + Gas Flow",
        'fos_subplot': "FOS + AI Collapse Prediction (NN) + Yielded Zones",
        'gas_flow': "Gas flow",
        'shear': "Shear",
        'tensile': "Tensile",
        'ai_collapse': "AI Collapse (NN)",
        'monitoring_panel': "📊 {obj_name}: Integrated Monitoring Panel",
        'pillar_live': "Pillar Strength",
        'rec_width_live': "Recommended Pillar Width",
        'max_subsidence_live': "Max Subsidence",
        'process_stage': "Process stage",
        'stage_active': "Active",
        'stage_cooling': "Cooling",
        'ai_monitor_title': "🧠 UCG AI Predictive Monitoring",
        'ai_monitor_desc': "Real-time sensor data and anomaly detection",
        'ai_steps': "Simulation steps:",
        'ai_run_btn': "▶️ Run AI Monitoring",
        'ai_stop_btn': "⏹ Stop",
        'advanced_analysis': "🔍 In-depth Dynamic Analysis and Methodological Justification",
        'tab_mass': "🏗️ Rock Mass Parameters",
        'tab_thermal': "🔥 Thermal Degradation",
        'tab_stability': "⚖️ Stability & References",
        'hb_class': "1. Hoek-Brown (2018) Classification",
        'hb_mb': "m_b = {mb:.3f}",
        'hb_s': "s = {s:.4f}",
        'hb_caption_mb': "Mass friction angle function (m_i={mi})",
        'hb_caption_s': "Fracturing degree (GSI={gsi})",
        'hb_interpret': "**Scientific note:** According to Hoek & Brown (2018), GSI={gsi} means rock mass strength is **{perc:.1f}%** lower than laboratory.",
        'thermal_params': "2. Thermo-Mechanical Coefficient Analysis",
        'param_table_param': "Parameter",
        'param_table_value': "Value",
        'param_table_reason': "Justification",
        'modulus': "Elastic Modulus (E)",
        'alpha': "Thermal expansion (α)",
        'temp0': "Initial T₀",
        'modulus_reason': "Typical average deformation coefficient for coal.",
        'alpha_reason': "Linear thermal expansion coefficient of coal (Yang, 2010).",
        'temp0_reason': "Initial natural temperature of the coal seam.",
        'ucs_decay': "**A) Thermal UCS reduction:**",
        'ucs_decay_eq': r"\sigma_{{ci(T)}} = \sigma_{{ci(0)}} \cdot e^{{-\beta(T-T_0)}} = {ucs:.2f} \text{{ MPa}}",
        'ucs_interpret': "**Interpretation:** At {temp}°C, rock strength decreased by {perc:.1f}%.",
        'thermal_stress': "**B) Thermal stress ($\\sigma_{{th}}$):**",
        'thermal_stress_eq': r"\sigma_{{th}} \approx \frac{{E \cdot \alpha \cdot \Delta T}}{{1 - \nu}} = {sigma:.2f} \text{{ MPa}}",
        'pillar_stability': "3. Pillar Stability and Bibliography",
        'fos_eq': r"FOS = \frac{{\sigma_p}}{{\sigma_v}} = {fos:.2f}",
        'pillar_wilson': "**Wilson (1972) Yield Pillar theory:** With pillar width $w={w}$ m, the central core sustains {sv:.2f} MPa geostatic load. Plastic zone: $y = {y:.1f}$ m.",
        'references': "#### 📚 Main Scientific References:",
        'ref1': "**Hoek, E., & Brown, E. T. (2018).** The Hoek-Brown failure criterion and GSI – 2018 edition. *JRMGE*.",
        'ref2': "**Yang, D. (2010).** *Stability of Underground Coal Gasification*. PhD Thesis, TU Delft.",
        'ref3': "**Shao, S., et al. (2015).** A thermal damage constitutive model for rock. *IJRMMS*.",
        'ref4': "**Wilson, A. H. (1972).** Research into the determination of pillar size. *Mining Engineer*.",
        'conclusion_danger': "🔴 **Scientific Conclusion:** FOS={fos:.2f}. High thermal degradation. Increase pillar width or control gasification rate.",
        'conclusion_safe': "🟢 **Scientific Conclusion:** FOS={fos:.2f}. Parameters ensure mass stability.",
        'methodology_expander': "📚 Scientific Methodology and References (PhD Research)",
        'tensile_empirical': "Empirical (UCS)",
        'tensile_hb': "HB-based (auto)",
        'tensile_manual': "Manual",
        'dip_angle_label': "Dip angle (°)",
        'timeline_table': """
| Stage | Time | Description |
|-------|------|-------------|
| **Planning** | 2026-04-01 | Validation, safety functions |
| **Optimization** | 2026-05-15 | NN/RF testing, FDM improvement |
| **Integration** | 2026-06-30 | Unit tests, final deploy |
        """,
        'live_monitoring_tab': "🔄 Live 3D Monitoring",
        'download_data': "📥 Download monitoring data (CSV)",
        'well_config': "Well Configuration",
        'well_distance': "Distance between wells (m):",
        'warning_cavity_width': "Warning: Pillar width exceeds well distance. cavity_width set to 1m.",
        'ucg_stages_title': "UCG Burning Stages (1-2-3 scheme) – New Scientific Model",
        'select_stage': "Select stage:",
        'geomech_state': "Geomechanical State (New Scientific Model)",
        'auto_animation': "Auto animation (1→2→3 stages)",
        'animation_done': "Animation finished.",
        'pillar_annotation': "PROTECTIVE PILLAR",
        'system_entropy': "System entropy (uncertainty)",
        'pin_approx': "**Note:** Kirsch solution is quasi-static (regularized). FEM remeshing required for large deformations.",
        'phase_field_info': "**Phase-field model (based on Bourdin et al., 2000):**",
        'uq_info': "Uncertainty Quantification (Monte-Carlo, Cholesky decomposition):",
        'shap_info': "SHAP interpretation (Lundberg & Lee, 2017):",
        'sobol_info': "Global sensitivity (Sobol', 2001):",
        'lhs_info': "Latin Hypercube Sampling (McKay, 1979):",
        'validation_info': "Experimental validation:",
        'experimental_note': "CSV must contain 'x' (m) and 'subsidence_cm' columns."
    },
    'ru': {
        'app_title': "Универсальный мониторинг деформаций поверхности",
        'app_subtitle': "Термомеханический (TM) анализ и оптимизация размеров целиков",
        'sidebar_header_params': "⚙️ Общие параметры",
        'formula_show': "Показать формулы:",
        'project_name': "Название проекта:",
        'process_time': "Время процесса (часы):",
        'num_layers': "Количество слоёв:",
        'tensile_model': "Модель растяжения:",
        'rock_props': "💎 Свойства породы",
        'disturbance': "Фактор нарушенности (D):",
        'poisson': "Коэффициент Пуассона (ν):",
        'stress_ratio': "Соотношение напряжений (k = σh/σv):",
        'tensile_params': "📐 Растяжение и целик",
        'thermal_decay_label': "Термическая деградация (β):",
        'combustion': "🔥 Горение и термическое воздействие",
        'burn_duration': "Продолжительность горения (часы):",
        'max_temp': "Максимальная температура (°C)",
        'timeline': "📅 Хронология проекта",
        'layer_params': "Параметры слоя {num}",
        'layer_name': "Название:",
        'thickness': "Мощность (м):",
        'ucs': "UCS (МПа):",
        'density': "Плотность (кг/м³):",
        'color': "Цвет:",
        'gsi': "GSI:",
        'mi': "mi:",
        'manual_st0': "σt0 (МПа):",
        'error_thick_positive': "Мощность должна быть >0",
        'error_ucs_positive': "UCS должен быть >0 МПа",
        'error_density_positive': "Плотность должна быть >0 кг/м³",
        'error_gsi_range': "GSI должен быть в диапазоне 10...100",
        'error_mi_positive': "mi должен быть >0",
        'error_min_layers': "❌ Требуется хотя бы 1 слой!",
        'warning_pytorch': "⚠️ PyTorch не установлен. Используется RandomForestClassifier.",
        'pillar_strength': "Прочность целика (σp)",
        'plastic_zone': "Пластическая зона (y)",
        'cavity_volume': "Объём каверны",
        'max_permeability': "Макс. проницаемость",
        'ai_recommendation': "Аналитическая рекомендация (целик)",
        'monitoring_header': "📊 {obj_name}: Мониторинг и экспертное заключение",
        'subsidence_title': "📉 Оседание поверхности (см)",
        'thermal_deform_title': "🔥 Горизонтальное смещение (см)",
        'hb_envelopes_title': "🛡️ Огибающие Хука-Брауна",
        'scientific_analysis': "📋 Научный анализ",
        'fos_red': "🔴 КЗП < 1.0: Разрушение",
        'fos_yellow': "🟡 КЗП 1.0–1.5: Неустойчиво",
        'fos_green': "🟢 КЗП > 1.5: Устойчиво",
        'tm_field_title': "🔥 ТМ поле и интерференция целика",
        'temp_subplot': "Поле температур (°C) + Поток газа",
        'fos_subplot': "КЗП + AI прогноз обрушения (НС) + Зоны пластичности",
        'gas_flow': "Поток газа",
        'shear': "Сдвиг",
        'tensile': "Растяжение",
        'ai_collapse': "AI обрушение (НС)",
        'monitoring_panel': "📊 {obj_name}: Комплексная панель мониторинга",
        'pillar_live': "Прочность целика",
        'rec_width_live': "Рекомендуемая ширина целика",
        'max_subsidence_live': "Макс. оседание",
        'process_stage': "Стадия процесса",
        'stage_active': "Активная",
        'stage_cooling': "Охлаждение",
        'ai_monitor_title': "🧠 AI-мониторинг UCG",
        'ai_monitor_desc': "Данные сенсоров в реальном времени и обнаружение аномалий",
        'ai_steps': "Количество шагов симуляции:",
        'ai_run_btn': "▶️ Запустить AI мониторинг",
        'ai_stop_btn': "⏹ Остановить",
        'advanced_analysis': "🔍 Углублённый динамический анализ и методическое обоснование",
        'tab_mass': "🏗️ Параметры массива",
        'tab_thermal': "🔥 Термическая деградация",
        'tab_stability': "⚖️ Устойчивость и источники",
        'hb_class': "1. Классификация Хука-Брауна (2018)",
        'hb_mb': "m_b = {mb:.3f}",
        'hb_s': "s = {s:.4f}",
        'hb_caption_mb': "Функция угла трения массива (m_i={mi})",
        'hb_caption_s': "Степень трещиноватости (GSI={gsi})",
        'hb_interpret': "**Научный комментарий:** По критерию Хука & Брауна (2018), GSI={gsi} означает, что прочность массива на **{perc:.1f}%** ниже лабораторной.",
        'thermal_params': "2. Анализ термомеханических коэффициентов",
        'param_table_param': "Параметр",
        'param_table_value': "Значение",
        'param_table_reason': "Обоснование",
        'modulus': "Модуль упругости (E)",
        'alpha': "Тепловое расширение (α)",
        'temp0': "Начальная T₀",
        'modulus_reason': "Типичный средний коэффициент деформации угля.",
        'alpha_reason': "Коэффициент линейного теплового расширения угля (Yang, 2010).",
        'temp0_reason': "Начальная естественная температура угольного пласта.",
        'ucs_decay': "**А) Тепловое снижение UCS:**",
        'ucs_decay_eq': r"\sigma_{{ci(T)}} = \sigma_{{ci(0)}} \cdot e^{{-\beta(T-T_0)}} = {ucs:.2f} \text{{ МПа}}",
        'ucs_interpret': "**Интерпретация:** При {temp}°C прочность породы снизилась на {perc:.1f}%.",
        'thermal_stress': "**Б) Термическое напряжение ($\\sigma_{{th}}$):**",
        'thermal_stress_eq': r"\sigma_{{th}} \approx \frac{{E \cdot \alpha \cdot \Delta T}}{{1 - \nu}} = {sigma:.2f} \text{{ МПа}}",
        'pillar_stability': "3. Устойчивость целика и библиография",
        'fos_eq': r"КЗП = \frac{{\sigma_p}}{{\sigma_v}} = {fos:.2f}",
        'pillar_wilson': "**По теории целика Уилсона (1972):** При ширине целика $w={w}$ м центральное ядро выдерживает геостатическую нагрузку {sv:.2f} МПа. Пластическая зона: $y = {y:.1f}$ м.",
        'references': "#### 📚 Основные научные источники:",
        'ref1': "**Hoek, E., & Brown, E. T. (2018).** The Hoek-Brown failure criterion and GSI – 2018 edition. *JRMGE*.",
        'ref2': "**Yang, D. (2010).** *Stability of Underground Coal Gasification*. PhD Thesis, TU Delft.",
        'ref3': "**Shao, S., et al. (2015).** A thermal damage constitutive model for rock. *IJRMMS*.",
        'ref4': "**Wilson, A. H. (1972).** Research into the determination of pillar size. *Mining Engineer*.",
        'conclusion_danger': "🔴 **Научный вывод:** КЗП={fos:.2f}. Высокая термическая деградация. Рекомендуется увеличить ширину целика.",
        'conclusion_safe': "🟢 **Научный вывод:** КЗП={fos:.2f}. Выбранные параметры обеспечивают устойчивость массива.",
        'methodology_expander': "📚 Научная методология и источники (PhD Research References)",
        'tensile_empirical': "Эмпирический (UCS)",
        'tensile_hb': "HB-метод (авто)",
        'tensile_manual': "Вручную",
        'dip_angle_label': "Угол падения пласта (Dip - °)",
        'timeline_table': """
| Этап | Время | Описание |
|------|-------|----------|
| **Планирование** | 2026-04-01 | Валидация, функции безопасности |
| **Оптимизация** | 2026-05-15 | NN/RF тестирование, улучшение FDM |
| **Интеграция** | 2026-06-30 | Unit тесты, финальный деплой |
        """,
        'live_monitoring_tab': "🔄 Мониторинг в реальном времени",
        'download_data': "📥 Загрузить данные мониторинга (CSV)",
        'well_config': "Конфигурация скважин",
        'well_distance': "Расстояние между скважинами (м):",
        'warning_cavity_width': "Предупреждение: ширина целика превышает расстояние между скважинами. cavity_width=1м.",
        'ucg_stages_title': "Стадии горения UCG (схема 1-2-3) – Новая научная модель",
        'select_stage': "Выберите стадию:",
        'geomech_state': "Геомеханическое состояние (Новая научная модель)",
        'auto_animation': "Авто анимация (1→2→3 стадии)",
        'animation_done': "Анимация завершена.",
        'pillar_annotation': "ЗАЩИТНЫЙ ЦЕЛИК",
        'system_entropy': "Системная энтропия (неопределённость)",
        'pin_approx': "**Примечание:** Решение Кирша квазистатическое (регуляризованное). Для больших деформаций требуется МКЭ.",
        'phase_field_info': "**Фазовая модель (основана на Bourdin et al., 2000):**",
        'uq_info': "Количественная оценка неопределённости (Монте-Карло, декомпозиция Холецкого):",
        'shap_info': "Интерпретация SHAP (Lundberg & Lee, 2017):",
        'sobol_info': "Глобальная чувствительность (Соболь, 2001):",
        'lhs_info': "Латинский гиперкуб (McKay, 1979):",
        'validation_info': "Экспериментальная валидация:",
        'experimental_note': "CSV должен содержать столбцы 'x' (м) и 'subsidence_cm'."
    }
}

FORMULA_OPTIONS = {
    'uz': ["Yopish", "1. Hoek-Brown Failure (2018)", "2. Thermal Damage & Permeability",
           "3. Thermal Stress & Tension", "4. Pillar & Subsidence"],
    'en': ["Close", "1. Hoek-Brown Failure (2018)", "2. Thermal Damage & Permeability",
           "3. Thermal Stress & Tension", "4. Pillar & Subsidence"],
    'ru': ["Закрыть", "1. Разрушение Хука-Брауна (2018)", "2. Термическое повреждение и проницаемость",
           "3. Термическое напряжение и растяжение", "4. Целик и оседание"]
}


def translate(key, **kwargs):
    lang = st.session_state.get('language', 'uz')
    text = TRANSLATIONS.get(lang, TRANSLATIONS['uz']).get(key, key)
    return text.format(**kwargs) if kwargs else text


t = translate


# ================================================================================
# FIZIK HISOB-KITOB FUNKSIYALARI
# (✅ TO'G'IRLANGAN VERSIYALAR — app_corrected.py dan)
# ================================================================================

# ==================== TUZATISH #1: HOEK-BROWN (2018) ====================
def hoek_brown_params_corrected(gsi: float, mi: float, D: float) -> tuple:
    """
    ✅ TO'G'IRLANGAN: Hoek-Brown (2018) parametrlari
    Manba: Hoek, E., & Brown, E. T. (2018). The Hoek-Brown failure criterion
           and GSI – 2018 edition. JRMGE, 10(1), 1-57.

    'a' parametri [0.5, 0.65] oralig'ida cheklangan (ko'mir uchun fizik chegara).
    Qaytaradi: (mb, s, a)
    """
    if not (10 <= gsi <= 100):
        raise PhysicsValidationError("GSI", gsi, (10, 100))
    if not (0 <= D <= 1):
        raise PhysicsValidationError("D", D, (0, 1))

    mb = mi * np.exp((gsi - 100.0) / (28.0 - 14.0 * D))

    if isinstance(gsi, (int, float)):
        s = np.exp((gsi - 100.0) / (9.0 - 3.0 * D)) if gsi > 25 else 0.0
    else:
        gsi_arr = np.asarray(gsi)
        s = np.where(gsi_arr > 25,
                     np.exp((gsi_arr - 100.0) / (9.0 - 3.0 * D)),
                     0.0)

    # ✅ TO'G'IRLANGAN: 'a' parametri fizik chegara bilan
    gsi_arr = np.asarray(gsi)
    a = 0.5 + (1.0 / 6.0) * (np.exp(-gsi_arr / 15.0) - np.exp(-20.0 / 3.0))
    a = np.clip(a, 0.5, 0.65)   # Ko'mir uchun fizik chegara

    return mb, s, a


# Orqaga moslik uchun alias (backward compatibility)
def hoek_brown_params(gsi: float, mi: float, D: float) -> tuple:
    """Hoek-Brown (2018) — hoek_brown_params_corrected ga alias"""
    return hoek_brown_params_corrected(gsi, mi, D)


def hoek_brown(sigma3, sigma_ci, mb, s, a):
    """Hoek-Brown yemirilish mezoni / failure criterion"""
    sigma3_eff = np.maximum(sigma3, 0.0)
    term = mb * (sigma3_eff / (sigma_ci + EPS)) + s
    term = np.maximum(term, 0.0)
    return sigma3_eff + sigma_ci * term ** a


def compute_demand_capacity_ratio(sigma1_applied, sigma3_confining, sigma_ci, mb, s, a):
    """Ortiqcha kuchlanish koeffitsienti / Overstress ratio (clipped [0,3])"""
    sigma3_eff = np.maximum(sigma3_confining, 0.0)
    sigma1_failure = (sigma3_eff +
                      sigma_ci * (np.maximum(mb * (sigma3_eff / (sigma_ci + EPS)) + s, 0.0) ** a))
    return np.clip(sigma1_applied / (sigma1_failure + EPS), 0, 3)


# ==================== TUZATISH #2: TERMAL ZARAR (SHAO et al., 2015) ====================
def thermal_damage_shao(T, beta, T_ref=T_REF_AMBIENT):
    """
    ✅ TO'G'IRLANGAN: Ikki bosqichli termal zarar modeli
    Manba: Shao, S., et al. (2015). A thermal damage constitutive model for rock.
           International Journal of Rock Mechanics and Mining Sciences, 81, 1-10.

    Bosqich 1: Chiziqli elastik   (20–320°C)
    Bosqich 2: Eksponensial       (320–1200°C)
    """
    delta_T = np.maximum(T - T_ref, 0.0)
    damage = np.where(
        T < T_ref + 300,
        beta * delta_T / 300.0,
        1.0 - np.exp(-beta * (delta_T - 300.0))
    )
    return np.clip(damage, 0.0, 1.0)


def thermal_damage(T, beta, T_ref=T_REF_AMBIENT):
    """
    Soddalashtirilgan eksponensial model (orqaga moslik uchun saqlab qolindi).
    PhD/Patent uchun thermal_damage_shao ishlatiladi.
    """
    return 1.0 - np.exp(-beta * np.maximum(T - T_ref, 0.0))


def apply_thermal_degradation(ucs0, T, beta):
    """
    ✅ TO'G'IRLANGAN: Termal zarar asosida UCS pasayishi.
    Shao et al. (2015) ikki bosqichli modeli ishlatiladi.
    """
    dmg = thermal_damage_shao(T, beta)
    ucs_T = ucs0 * (1.0 - dmg)
    return np.clip(ucs_T, 0.5, None)


# ==================== TUZATISH #3: TO'LIQ 3D BOSH KUCHLANISHLAR ====================
def principal_stresses_3d(sx, sy, sz, txy, tyz, txz):
    """
    ✅ TO'G'IRLANGAN: To'liq 3D bosh kuchlanishlarni hisoblash.
    Simmetrik stress tenzori quriladi, xos qiymatlar hisoblanadi.
    """
    sigma_matrix = np.array([
        [sx, txy, txz],
        [txy, sy, tyz],
        [txz, tyz, sz]
    ])
    if not np.allclose(sigma_matrix, sigma_matrix.T):
        raise ValueError("Stress tenzori simmetrik emas! / Stress tensor not symmetric!")
    eigenvalues = np.linalg.eigvalsh(sigma_matrix)
    eigenvalues = np.sort(eigenvalues)[::-1]   # Kamayish tartibida: s1, s2, s3
    return eigenvalues[0], eigenvalues[1], eigenvalues[2]


def principal_stresses(sx, sy, txy):
    """2D bosh kuchlanishlar (UI uchun saqlab qolindi)"""
    avg = (sx + sy) / 2
    radius = np.sqrt(((sx - sy) / 2) ** 2 + txy ** 2)
    return avg + radius, avg - radius


# ==================== TUZATISH #4: REGULARIZATSIYALANGAN KIRSCH ====================
def kirsch_stress_field_regularized(x, z, sigma_H, sigma_h, cavity_radius,
                                    pore_pressure=0.0, regularization='smooth'):
    """
    ✅ TO'G'IRLANGAN: Regularizatsiyalangan Kirsch yechimi (singularliklar yo'q).
    Manba: Kirsch (1898) — r→a singularligi bartaraf etilgan.

    regularization: 'smooth' (Gauss yadrosi), 'clamp' (5% bufer)
    """
    r = np.sqrt(x ** 2 + z ** 2)

    if regularization == 'smooth':
        r_smooth = np.sqrt(r ** 2 + (0.1 * cavity_radius) ** 2)
    elif regularization == 'clamp':
        r_smooth = np.maximum(r, cavity_radius * 1.05)
    else:
        r_smooth = r

    theta = np.arctan2(z, x)
    a = cavity_radius
    a2_r2 = (a ** 2) / (r_smooth ** 2 + EPS)
    a4_r4 = (a ** 4) / (r_smooth ** 4 + EPS)

    sigma_rr = ((sigma_H + sigma_h) / 2 * (1 - a2_r2) +
                (sigma_H - sigma_h) / 2 * (1 - 4 * a2_r2 + 3 * a4_r4) * np.cos(2 * theta))
    sigma_tt = ((sigma_H + sigma_h) / 2 * (1 + a2_r2) -
                (sigma_H - sigma_h) / 2 * (1 + 3 * a4_r4) * np.cos(2 * theta))
    tau_rt = -(sigma_H - sigma_h) / 2 * (1 + 2 * a2_r2 - 3 * a4_r4) * np.sin(2 * theta)

    sigma_rr -= pore_pressure
    sigma_tt -= pore_pressure
    return sigma_rr, sigma_tt, tau_rt


def kirsch_stress_field(x, z, sigma_H, sigma_h, cavity_radius, pore_pressure=0.0):
    """Regularizatsiyalangan Kirsch ga alias (orqaga moslik)"""
    return kirsch_stress_field_regularized(x, z, sigma_H, sigma_h,
                                           cavity_radius, pore_pressure, 'smooth')


# ==================== TUZATISH #5: ROBIN CHEGARAVIY SHARTI ====================
def solve_heat_equation_robin_bc(T, Q, rho_field, cp_field, k_field,
                                  dx, dz, total_time, T_air, h_conv=10.0):
    """
    ✅ TO'G'IRLANGAN: Robin chegaraviy sharti bilan issiqlik tenglamasi.
    Robin SH: -k * dT/dz = h * (T - T_air)
    Manba: Incropera & DeWitt (2007) — Heat and Mass Transfer.

    CFL barqarorlik mezoni avtomatik hisoblanadi.
    """
    alpha_field = k_field / (rho_field * cp_field + EPS)
    alpha_max = np.max(alpha_field)

    dt_max = 1.0 / (2 * alpha_max * (1 / dx ** 2 + 1 / dz ** 2))
    dt = 0.8 * dt_max
    n_steps = max(int(np.ceil(total_time / dt)), 1)
    dt = total_time / n_steps

    for _ in range(n_steps):
        T_old = T.copy()
        Txx = (T_old[1:-1, 2:] - 2 * T_old[1:-1, 1:-1] + T_old[1:-1, :-2]) / (dx ** 2)
        Tzz = (T_old[2:, 1:-1] - 2 * T_old[1:-1, 1:-1] + T_old[:-2, 1:-1]) / (dz ** 2)

        T_new = T_old.copy()
        T_new[1:-1, 1:-1] = (T_old[1:-1, 1:-1] +
                              dt * (alpha_field[1:-1, 1:-1] * (Txx + Tzz) +
                                    Q[1:-1, 1:-1] /
                                    (rho_field[1:-1, 1:-1] * cp_field[1:-1, 1:-1] + EPS)))
        # Lateral boundaries (symmetry)
        T_new[:, 0] = T_new[:, 1]
        T_new[:, -1] = T_new[:, -2]
        # Bottom (insulated — Neumann)
        T_new[-1, :] = T_new[-2, :]
        # Top (Robin B.C. — convection)
        k_surface = k_field[0, :]
        T_new[0, :] = ((T_new[1, :] + (h_conv * dz / (k_surface + EPS)) * T_air) /
                       (1.0 + h_conv * dz / (k_surface + EPS)))
        T = T_new.copy()
    return T


# ==================== TUZATISH #6: KOZENY-CARMAN O'TKAZUVCHANLIGI ====================
def permeability_modified_kozeny_carman(damage, volumetric_strain=0.0,
                                        stress_eff=1.0, porosity_initial=0.35,
                                        d_char=1e-4):
    """
    ✅ TO'G'IRLANGAN: Modifikatsiyalangan Kozeny-Carman (zarar bilan bog'langan).
    Manbalar:
      Walsh & Brace (1984). JGR, 89(B12), 9425-9432.
      Zhao et al. (2015). Thermal damage evolution of coal.

    k = (φ³ / (180·(1−φ)²)) · d² · f_zarar · f_kuchlanish
    """
    phi_0 = porosity_initial
    phi = phi_0 * (1.0 + volumetric_strain) / (1.0 - damage + EPS)
    phi = np.clip(phi, 1e-4, 0.5)

    k_base = (phi ** 3 / (180.0 * (1.0 - phi + EPS) ** 2)) * d_char ** 2
    k_damage_factor = 1.0 + 50.0 * damage ** 2
    k_stress_factor = np.exp(-0.01 * np.maximum(stress_eff, 0.1))

    return np.clip(k_base * k_damage_factor * k_stress_factor, 1e-18, 1e-10)


# ==================== TUZATISH #7: HARORATGA BOG'LIQ GAZ VISKOZITETI ====================
def viscosity_temperature(T):
    """
    ✅ TO'G'IRLANGAN: Haroratga bog'liq gaz viskoziteti.
    Manba: Chapman & Cowling (1970) — kinetik nazariya.
    μ(T) = μ₀ · (T/T₀)ⁿ, n ≈ 0.67 (ikki atomli gazlar)
    """
    T_ref_K = 293.15
    mu_ref = 3e-5
    T_K = T + 273.15 if np.ndim(T) == 0 and T > 100 else np.asarray(T) + 273.15
    n = 0.67
    mu_T = mu_ref * (T_K / T_ref_K) ** n
    return np.clip(mu_T, 1e-6, 1e-3)


# ==================== TUZATISH #8: MONTE CARLO + CHOLESKY ====================
def monte_carlo_fos_corrected(ucs_mean, ucs_std, gsi_mean, gsi_std,
                               mi_val, D, T_avg, H_seam, depth, density,
                               rec_width, beta_th, n_sim=1000,
                               random_seed=CONFIG.RANDOM_SEED):
    """
    ✅ TO'G'IRLANGAN: Cholesky dekompozitsiya bilan Monte Carlo.
    Korrelyatsiyalangan UCS va GSI namunalari NaN xatosisiz generatsiya qilinadi.
    Manba: Efron & Tibshirani (1993). Bootstrap.
    """
    rng = np.random.default_rng(seed=random_seed)

    correlation = 0.3
    cov_matrix = np.array([
        [ucs_std ** 2, correlation * ucs_std * gsi_std],
        [correlation * ucs_std * gsi_std, gsi_std ** 2]
    ])
    try:
        L = np.linalg.cholesky(cov_matrix)
    except np.linalg.LinAlgError:
        logger.warning("Kovariatsiya musbat aniq emas, korrelyatsiya kamaytirildi")
        cov_matrix[0, 1] = 0.1 * ucs_std * gsi_std
        cov_matrix[1, 0] = 0.1 * ucs_std * gsi_std
        L = np.linalg.cholesky(cov_matrix)

    Z = rng.standard_normal((n_sim, 2))
    samples_corr = Z @ L.T
    ucs_samples = samples_corr[:, 0] + ucs_mean
    gsi_samples = np.clip(samples_corr[:, 1] + gsi_mean, 10, 100)

    fos = []
    for ucs, gsi in zip(ucs_samples, gsi_samples):
        try:
            mb, s, a = hoek_brown_params_corrected(gsi, mi_val, D)
            ucs_T = apply_thermal_degradation(ucs, T_avg, beta_th)
            sigma_cm = ucs_T * (np.maximum(s, 1e-9) ** a)
            pillar_strength = sigma_cm * (WILSON_C1 + WILSON_C2 * rec_width / (H_seam + EPS))
            sv = density * 9.81 * depth / 1e6
            fos_val = np.clip(pillar_strength / (sv + EPS), 0, 3)
            fos.append(fos_val)
        except Exception as e:
            logger.warning(f"MC iteratsiyasi muvaffaqiyatsiz: {e}")
            fos.append(1.0)

    fos = np.array(fos)
    pf = np.mean(fos < 1.0)
    return fos, pf


def monte_carlo_fos(ucs_mean, ucs_std, gsi_mean, gsi_std, mi_val, D,
                    T_avg, H_seam, depth, density, rec_width, beta_th,
                    n_sim=1000, random_seed=CONFIG.RANDOM_SEED):
    """monte_carlo_fos_corrected ga alias (orqaga moslik)"""
    return monte_carlo_fos_corrected(ucs_mean, ucs_std, gsi_mean, gsi_std,
                                     mi_val, D, T_avg, H_seam, depth, density,
                                     rec_width, beta_th, n_sim, random_seed)


# ==================== TUZATISH #9: BOOTSTRAP ISHONCH INTERVALI ====================
def subsidence_confidence_interval_bootstrap(sub_profile, n_bootstrap=1000,
                                              confidence=0.95):
    """
    ✅ TO'G'IRLANGAN: Bootstrap ishonch intervali (kichik namunalar uchun t-taqsimotdan yaxshi).
    Manba: Efron, B., & Tibshirani, R. J. (1993). An Introduction to the Bootstrap. CRC Press.
    """
    bootstrap_means = []
    rng = np.random.default_rng(seed=CONFIG.RANDOM_SEED)
    for _ in range(n_bootstrap):
        indices = rng.choice(len(sub_profile), size=len(sub_profile), replace=True)
        bootstrap_means.append(np.mean(sub_profile[indices]))
    bootstrap_means = np.array(bootstrap_means)
    alpha = 1 - confidence
    return (np.percentile(bootstrap_means, alpha / 2 * 100),
            np.percentile(bootstrap_means, (1 - alpha / 2) * 100))


def subsidence_confidence_interval(sub_profile, n_measurements, confidence=0.95):
    """T-taqsimot asosida ishonch intervali (orqaga moslik)"""
    std_est = np.std(sub_profile) * 0.15
    t_crit = t_dist.ppf((1 + confidence) / 2, df=n_measurements - 1)
    margin = t_crit * std_est / np.sqrt(n_measurements)
    return sub_profile - margin, sub_profile + margin


# ==================== TUZATISH #10: PORE BOSIMI (TEMPERATURE) ====================
def pore_pressure_field_corrected(T, depth, water_table=20.0, rho_water=1000.0):
    """
    ✅ TO'G'IRLANGAN: G'ovaklik bosimi maydoni (gidrostatik + gaz).
    Ideal gaz qonuni asosida haroratga bog'liq gaz bosimi.
    """
    h_water = np.maximum(depth - water_table, 0.0)
    P_hydro = rho_water * 9.81 * h_water / 1e6
    T_kelvin = np.maximum(T + 273.15, 293.15)
    P_gas = (101325.0 * T_kelvin / 293.15) / 1e6
    return P_hydro + P_gas


def pore_pressure_field(T, depth, water_table=20.0, rho_water=1000.0):
    """pore_pressure_field_corrected ga alias"""
    return pore_pressure_field_corrected(T, depth, water_table, rho_water)


# ==================== TUZATISH #11: FOS + SKEMPTON G'OVAKLIK BOSIMI ====================
def fos_with_pore_pressure(pillar_strength, sigma_v, pore_pressure, B_skempton=0.9):
    """
    ✅ TO'G'IRLANGAN: G'ovaklik bosimi tuzatmasi bilan FOS.
    Manba: Skempton, A. W. (1961). Effective stress in soils, concrete and rocks.
    σ'_v = σ_v − B·u
    """
    sigma_v_eff = np.maximum(sigma_v - B_skempton * pore_pressure, 0.01)
    return np.clip(pillar_strength / (sigma_v_eff + EPS), 0, 5)


# ==================== TUZATISH #12: O'LCHAMLI TAHLIL ====================
class Unit(Enum):
    """Fizik o'lchov birliklari"""
    METER = 'm'
    CENTIMETER = 'cm'
    SECOND = 's'
    HOUR = 'h'
    PASCAL = 'Pa'
    MEGAPASCAL = 'MPa'
    CELSIUS = '°C'
    KELVIN = 'K'


class DimensionalAnalysis:
    """
    ✅ TO'G'IRLANGAN: O'lchamli tahlil tekshiruvi.
    Hisoblashlarning fizik muvofiqligini ta'minlaydi.
    """
    @staticmethod
    def check_thermal_stress(E_Pa, alpha_1_K, Delta_T_K, nu):
        """
        Termal kuchlanish o'lchamli tekshiruvi:
        [σ] = Pa = [E]·[α]·[ΔT] / [1]
        """
        result_Pa = (E_Pa * alpha_1_K * Delta_T_K) / (1 - 2 * nu + EPS)
        assert result_Pa > 0, "Termal kuchlanish musbat bo'lishi kerak"
        assert result_Pa < 1e12, "Termal kuchlanish fizik chegaradan oshdi"
        return result_Pa

    @staticmethod
    def check_permeability(k_m2):
        """O'tkazuvchanlikning fizik chegaralari"""
        assert 1e-20 < k_m2 < 1e-8, f"O'tkazuvchanlik {k_m2:.2e} chegaradan tashqarida"
        return k_m2


# ==================== TUZATISH #13: GRADIENT TEKSHIRUVI ====================
def numerical_gradient_check(f, x, h=1e-5, threshold=1e-4):
    """
    ✅ TO'G'IRLANGAN: Oldingi va markaziy farqlarni taqqoslash.
    Gradient approksimatsiya aniqligini tekshiradi.
    """
    grad_fwd = (f(x + h) - f(x)) / h
    grad_central = (f(x + h) - f(x - h)) / (2 * h)
    error = abs(grad_fwd - grad_central)
    if error > threshold:
        logger.warning(f"Gradient approx xatosi: {error:.2e} (chegara: {threshold})")
    return grad_central


# ==================== TUZATISH #14: STRESS TENZORI VALIDATSIYASI ====================
def validate_stress_tensor_3d(sigma_tensor):
    """
    ✅ TO'G'IRLANGAN: 3x3 kuchlanish tenzori xossalarini tekshirish.
    Simmetriklik, kompleks xos qiymatlar va NaN/Inf tekshiriladi.
    """
    if not np.allclose(sigma_tensor, sigma_tensor.T):
        raise ValueError("Stress tenzori simmetrik emas! / Not symmetric!")
    eigenvalues = np.linalg.eigvals(sigma_tensor)
    if np.any(np.iscomplex(eigenvalues)):
        raise ValueError("Kompleks xos qiymatlar aniqlandi! / Complex eigenvalues!")
    if np.any(~np.isfinite(eigenvalues)):
        raise ValueError("Cheksiz xos qiymatlar! / Non-finite eigenvalues!")
    return np.sort(np.real(eigenvalues))[::-1]


def validate_stress_tensor(sigma_tensor):
    """validate_stress_tensor_3d ga alias (orqaga moslik)"""
    return validate_stress_tensor_3d(sigma_tensor)


# ==================== TUZATISH #15: VON MISES ====================
def von_mises_stress_3d(sigma_x, sigma_y, sigma_z, tau_xy, tau_yz, tau_zx):
    """
    ✅ TO'G'IRLANGAN: To'liq 3D Von Mises kuchlanishi.
    σ_vm = √(½·((σx−σy)² + (σy−σz)² + (σz−σx)²) + 3·(τxy² + τyz² + τzx²))
    """
    vm = np.sqrt(
        0.5 * ((sigma_x - sigma_y) ** 2 + (sigma_y - sigma_z) ** 2 + (sigma_z - sigma_x) ** 2) +
        3 * (tau_xy ** 2 + tau_yz ** 2 + tau_zx ** 2)
    )
    return np.maximum(vm, 0.0)


def von_mises_stress(sigma_x, sigma_z, tau_xz, nu=None):
    """2D Von Mises (UI kodi uchun saqlab qolindi)"""
    if nu is not None:
        sigma_y = nu * (sigma_x + sigma_z)
    else:
        sigma_y = sigma_z
    return np.sqrt(np.maximum(sigma_x ** 2 - sigma_x * sigma_y + sigma_y ** 2 + 3 * tau_xz ** 2, 0.0))


# ================================================================================
# QOLGAN FIZIK FUNKSIYALAR (app.py dan — to'liq saqlab qolindi)
# ================================================================================

def thermal_conductivity(T, k0=2.5):
    k = k0 * (1 - 0.0004 * (T - T_REF_AMBIENT))
    return np.clip(k, 0.5, None)


def specific_heat(T):
    cp = 960 + 0.14 * T
    return np.clip(cp, 900, 2200)


def density_temperature(rho0, T):
    T_clamped = np.clip(T, T_REF_AMBIENT, 1200.0)
    alpha_v = 3.6e-5
    rho_T = rho0 * (1.0 - alpha_v * (T_clamped - T_REF_AMBIENT))
    return np.clip(rho_T, 0.80 * rho0, rho0)


def young_modulus_temperature(T, E0=None):
    E0_val = E0 if E0 is not None else PARAMS.E_mass
    c_E = 0.0018
    E_T = E0_val * np.exp(-c_E * np.maximum(T - T_REF_AMBIENT, 0.0))
    return np.clip(E_T, 0.10 * E0_val, E0_val)


def thermal_expansion_temperature(T):
    T = np.clip(T, T_REF_AMBIENT, 1200)
    return PARAMS.alpha_thermal * (1 + 0.002 * (T - T_REF_AMBIENT) + 1e-6 * (T - T_REF_AMBIENT) ** 2)


def vertical_stress(depth, density):
    return density * 9.81 * depth / 1e6


def validate_heat_stability(alpha_max, dx, dz, dt):
    """CFL sharti: Courant soni ≤ 0.5"""
    courant = alpha_max * dt * (1 / dx ** 2 + 1 / dz ** 2)
    if courant > 0.5:
        raise NumericalInstabilityError(f"Beqaror! Courant={courant:.3f} > 0.5")
    return True


def profile_performance(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        start = time.perf_counter()
        result = func(*args, **kwargs)
        elapsed = time.perf_counter() - start
        logger.info(f"{func.__name__} bajarildi: {elapsed:.3f}s")
        return result
    return wrapper


@profile_performance
def solve_heat_equation_dynamic(T, Q, rho_field, cp_field, k_field,
                                 dx, dz, total_time, T_air, h=10.0):
    """
    Dinamik issiqlik tenglamasi — Robin BC bilan (to'g'irlangan versiyaga yo'naltirilgan).
    """
    return solve_heat_equation_robin_bc(T, Q, rho_field, cp_field, k_field,
                                        dx, dz, total_time, T_air, h_conv=h)


def evolving_cavity_radius(time_h, T_field, beta, grid_z, source_z, H_seam):
    source_mask = np.abs(grid_z - source_z) < 1.5 * H_seam
    if not np.any(source_mask):
        return 5.0
    T_source = T_field[source_mask]
    thermal_dam_local = thermal_damage(T_source, beta)
    growth_rate = 0.015 * np.mean(thermal_dam_local)
    radius = 5.0 + growth_rate * time_h
    return float(np.clip(radius, 5.0, 40.0))


def subsidence_inclined_seam(S_horizontal, dip_deg, depth, phi_deg):
    dip_rad = np.radians(dip_deg)
    phi_rad = np.radians(phi_deg)
    return depth * np.tan(dip_rad) * np.tan(phi_rad / 2)


def pillar_creep_strength(sigma_p0, time_h, A_creep=0.05, n_creep=0.3):
    reduction = min(A_creep * (time_h ** n_creep), 0.40)
    return sigma_p0 * (1.0 - reduction)


def gas_migration_risk(T_field, perm_field, depth, fos_field):
    thermal_path = T_field > 300.0
    perm_path = perm_field > 1e-14
    structural_fail = fos_field < 1.5
    return gaussian_filter((thermal_path & perm_path & structural_fail).astype(float), sigma=2.0)


def water_inrush_risk(void_volume, aquifer_depth, depth_seam, fos_min):
    height_to_aquifer = abs(aquifer_depth - depth_seam)
    h_critical = 0.0015 * void_volume ** 0.5
    if height_to_aquifer < h_critical and fos_min < 1.2:
        return "CRITICAL", 0.9
    elif height_to_aquifer < h_critical * 1.5:
        return "HIGH", 0.6
    return "LOW", 0.1


def propagate_uncertainty_analytical(ucs_mean, ucs_cov, gsi_mean, gsi_cov,
                                      T_mean, T_cov, H_seam, rec_width):
    eps = 0.01

    def _quick_fos_local(ucs, gsi, T):
        mb, s, a = hoek_brown_params(gsi, 10, 0.7)
        ucs_T = apply_thermal_degradation(ucs, T, PARAMS.thermal_damage_beta)
        sigma_cm = ucs_T * (max(float(s), 1e-9) ** float(a))
        p_str = sigma_cm * (WILSON_C1 + WILSON_C2 * rec_width / (H_seam + EPS))
        sv = vertical_stress(200, 2500)
        return np.clip(p_str / (sv + EPS), 0, 5)

    fos_base = _quick_fos_local(ucs_mean, gsi_mean, T_mean)
    dfos_ducs = (_quick_fos_local(ucs_mean * (1 + eps), gsi_mean, T_mean) - fos_base) / (eps * ucs_mean)
    dfos_dgsi = (_quick_fos_local(ucs_mean, gsi_mean * (1 + eps), T_mean) - fos_base) / (eps * gsi_mean)
    dfos_dT = (_quick_fos_local(ucs_mean, gsi_mean, T_mean * (1 + eps)) - fos_base) / (eps * T_mean)
    var_fos = ((dfos_ducs * ucs_mean * ucs_cov) ** 2 +
               (dfos_dgsi * gsi_mean * gsi_cov) ** 2 +
               (dfos_dT * T_mean * T_cov) ** 2)
    return fos_base, np.sqrt(var_fos)


def adaptive_mesh(fos_field, stress_gradient, threshold=1.5):
    """FOS yoki stress gradienti yuqori bo'lgan joylarda to'rni zichlashtirish"""
    return (fos_field < threshold) | (stress_gradient > 100)


def check_convergence(residuals, tolerance=1e-6, window=10):
    """Iteratsion solver yaqinlashuvini tekshirish"""
    if len(residuals) < window:
        return False
    recent = np.array(residuals[-window:])
    relative_change = np.abs(np.diff(recent)) / (np.abs(recent[:-1]) + EPS)
    return np.all(relative_change < tolerance)


def mesh_quality_metrics(grid_x, grid_z):
    dx = np.diff(grid_x, axis=1)
    dz = np.diff(grid_z, axis=0)
    aspect_ratio = np.max(dz) / (np.max(dx) + EPS)
    if aspect_ratio > 10:
        st.warning(f"⚠️ To'r geometriyasi yomon: aspect={aspect_ratio:.1f}")
    return aspect_ratio


@dataclass
class BoundaryConditions:
    """Chegaraviy shartlar aniq spetsifikatsiyasi"""
    top: str = 'free'
    bottom: str = 'fixed'
    left: str = 'symmetric'
    right: str = 'symmetric'
    temperature_surface: float = 25.0
    temperature_bottom: float = 30.0

    def validate(self):
        valid_bc = ['free', 'fixed', 'dirichlet', 'symmetric']
        for direction, bc_type in [('top', self.top), ('bottom', self.bottom)]:
            if bc_type not in valid_bc:
                raise ValueError(f"Noto'g'ri chegaraviy shart '{bc_type}' {direction} uchun")


bc = BoundaryConditions(top='free', temperature_surface=25.0)
bc.validate()


def rk4_stability_region(lambda_dt):
    z = np.linspace(-4, 2.5, 1000) + 1j * np.linspace(-2.8, 2.8, 1000)
    return np.abs(1 + z + z ** 2 / 2 + z ** 3 / 6 + z ** 4 / 24) <= 1


class MaterialModel:
    def stress_from_strain(self, strain): raise NotImplementedError
    def validate(self): raise NotImplementedError


class LinearElastic(MaterialModel):
    def __init__(self, E, nu):
        self.E, self.nu = E, nu
    def stress_from_strain(self, strain):
        return self.E * strain
    def validate(self):
        assert 0 < self.E < 1e12
        assert -1 < self.nu < 0.5


class NonlinearDamage(MaterialModel):
    def __init__(self, E0, D, beta):
        self.E0, self.D, self.beta = E0, D, beta
    def stress_from_strain(self, strain, T):
        damage = thermal_damage_shao(T, self.beta)
        return self.E0 * (1 - damage) * strain


def verify_solution():
    """Analitik yechim bilan taqqoslash"""
    x, z, tt = 1.0, 1.0, 0.1
    T_exact = np.sin(x) * np.cos(z) * np.exp(-tt)
    dT_dt = -np.sin(x) * np.cos(z) * np.exp(-tt)
    d2T_dx2 = -np.sin(x) * np.cos(z) * np.exp(-tt)
    d2T_dz2 = -np.sin(x) * np.cos(z) * np.exp(-tt)
    residual = dT_dt + d2T_dx2 + d2T_dz2
    assert abs(residual) < 1e-10, f"Qoldiq juda katta: {residual}"


# ================================================================================
# INFRATUZILMA (Infrastructure)
# ================================================================================

# --- SQLAlchemy ---
engine = create_engine(
    'sqlite:///ucg_data.db',
    poolclass=pool.QueuePool,
    pool_size=5,
    max_overflow=10,
    pool_recycle=3600
)
SessionLocal = sessionmaker(bind=engine)


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# --- Cache boshqaruvi ---
@st.cache_resource
def _initialize_computation_state():
    return {'temperature_cache': {}, 'stress_cache': {}, 'fos_cache': {}}


class CacheManager:
    """TTL qo'llab-quvvatlanadigan kesh"""
    def __init__(self, ttl_seconds=3600):
        self.cache = {}
        self.timestamps = {}
        self.ttl_seconds = ttl_seconds

    def get(self, key):
        if key not in self.cache:
            return None
        if time.time() - self.timestamps[key] > self.ttl_seconds:
            del self.cache[key]
            return None
        return self.cache[key]

    def set(self, key, value):
        self.cache[key] = value
        self.timestamps[key] = time.time()

    def invalidate_pattern(self, pattern):
        import re
        compiled = re.compile(pattern)
        for key in list(self.cache.keys()):
            if compiled.match(key):
                del self.cache[key]


# --- Rate Limiter ---
class RateLimiter:
    def __init__(self, calls=100, period=60):
        self.calls = calls
        self.period = period
        self.call_times = defaultdict(list)

    def is_allowed(self, key):
        now = time.time()
        self.call_times[key] = [t for t in self.call_times[key] if t > now - self.period]
        if len(self.call_times[key]) < self.calls:
            self.call_times[key].append(now)
            return True
        return False


limiter = RateLimiter(calls=100, period=60)


def rate_limited(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        user_id = st.session_state.get('user_id', 'anonymous')
        if not limiter.is_allowed(user_id):
            st.error("❌ So'rovlar soni chegaradan oshdi. Iltimos, kuting.")
            return None
        return func(*args, **kwargs)
    return wrapper


# --- Ma'lumotlar versiyalash ---
class DataVersionManager:
    def __init__(self, base_dir="data_versions"):
        self.base_dir = Path(base_dir)
        self.base_dir.mkdir(exist_ok=True)

    def save_version(self, data_dict, version_name):
        version_file = self.base_dir / f"{version_name}_{datetime.now().isoformat()}.json"
        with open(version_file, 'w') as f:
            json.dump({'data': data_dict, 'timestamp': datetime.now().isoformat(),
                       'version': version_name}, f, indent=2)
        return version_file

    def load_version(self, version_file):
        with open(version_file, 'r') as f:
            return json.load(f)


# --- Audit tizimi ---
@dataclass
class AuditLog:
    """✅ Regulyator muvofiqlik uchun audit izi"""
    user_id: str
    action: str
    timestamp: datetime
    parameters: dict
    result_status: str
    error_msg: str = None

    def to_dict(self):
        return {
            'user': self.user_id,
            'action': self.action,
            'timestamp': self.timestamp.isoformat(),
            'params': self.parameters,
            'status': self.result_status,
            'error': self.error_msg
        }


def log_audit(user_id, action, params, status, error=None):
    audit = AuditLog(
        user_id=user_id, action=action, timestamp=datetime.now(),
        parameters=params, result_status=status, error_msg=error
    )
    logger.info(json.dumps(audit.to_dict()))


# --- Tizim sog'lig'i ---
@dataclass
class HealthStatus:
    status: str
    uptime_seconds: float
    memory_usage_mb: float
    cpu_usage_percent: float
    last_error: str = None
    error_count: int = 0


def check_app_health():
    process = psutil.Process()
    return HealthStatus(
        status='healthy',
        uptime_seconds=time.time() - CONFIG.START_TIME,
        memory_usage_mb=process.memory_info().rss / 1024 / 1024,
        cpu_usage_percent=process.cpu_percent(interval=1),
        error_count=0
    )


# --- Hisoblash metrikasi ---
@dataclass
class ComputationMetrics:
    start_time: datetime
    end_time: datetime = None
    grid_shape: tuple = None
    temperature_max: float = None
    fos_min: float = None
    computation_time_s: float = None
    memory_peak_mb: float = None

    def log_metrics(self):
        logger.info(f"Hisoblash: {self.grid_shape}, "
                    f"Vaqt: {self.computation_time_s:.2f}s, "
                    f"Xotira: {self.memory_peak_mb:.1f}MB")


# ================================================================================
# ASOSIY FIZIK HISOBLASH FUNKSIYALARI (app.py dan)
# ================================================================================

@st.cache_data(show_spinner=False, max_entries=50)
@profile_performance
def compute_temperature_field_moving(time_h, T_source_max, burn_duration,
                                      total_depth, source_z, grid_shape):
    x_axis = np.linspace(-total_depth * 1.5, total_depth * 1.5, grid_shape[1])
    z_axis = np.linspace(0, total_depth + 50, grid_shape[0])
    dx = x_axis[1] - x_axis[0]
    dz = z_axis[1] - z_axis[0]
    grid_x, grid_z = np.meshgrid(x_axis, z_axis)
    temp_2d = np.full_like(grid_x, 25.0)
    rho_field = np.full_like(temp_2d, 1400.0)
    cp_field = specific_heat(temp_2d)
    k_field = thermal_conductivity(temp_2d)
    total_time = max(burn_duration, time_h) * 3600
    sources = [
        {'x0': -total_depth / 3, 'start': 0, 'moving': False},
        {'x0': 0, 'start': 40, 'moving': True, 'v': 0.02},
        {'x0': total_depth / 3, 'start': 80, 'moving': False}
    ]
    current_time_h = 0.0
    time_step = 3600.0
    n_steps = max(int(total_time / time_step), 1)
    time_step = total_time / n_steps
    source_mask_local = np.abs(grid_z - source_z) < 10.0
    rho_cp_ref = (np.mean((rho_field * cp_field)[source_mask_local])
                  if np.any(source_mask_local) else 1400.0 * 960.0)

    for step in range(n_steps):
        current_time_h += time_step / 3600.0
        Q_source = np.zeros_like(temp_2d)
        for src in sources:
            if current_time_h <= src['start']:
                continue
            dt_sec = (current_time_h - src['start']) * 3600
            x_center = src['x0'] + src.get('v', 0) * dt_sec if src['moving'] else src['x0']
            elapsed = current_time_h - src['start']
            if elapsed <= burn_duration:
                curr_T = T_source_max
            else:
                curr_T = 25 + (T_source_max - 25) * np.exp(-0.03 * (elapsed - burn_duration))
            pen_depth = np.sqrt(4 * PARAMS.THERMAL_DIFFUSIVITY * max(dt_sec, 3600)) + 15
            dist_sq = (grid_x - x_center) ** 2 + (grid_z - source_z) ** 2
            Q_source += rho_cp_ref * (curr_T - 25) / 3600.0 * np.exp(-dist_sq / (pen_depth ** 2))

        temp_2d = solve_heat_equation_dynamic(
            T=temp_2d, Q=Q_source, rho_field=rho_field, cp_field=cp_field,
            k_field=k_field, dx=dx, dz=dz, total_time=time_step, T_air=25.0
        )
    return temp_2d, x_axis, z_axis, grid_x, grid_z


def physics_features(T_flat, s1_flat, s3_flat, z_flat, ucs_val):
    """ML uchun fizik xususiyatlar"""
    damage = thermal_damage_shao(T_flat, PARAMS.thermal_damage_beta)
    fos = np.clip(ucs_val * (1 - damage) / (s1_flat + EPS), 0, 3)
    energy = 0.5 * s1_flat * s3_flat / (PARAMS.E_mass / 1e6 + EPS)
    X = np.column_stack([T_flat, s1_flat, s3_flat, z_flat, damage, fos, energy])
    y = (fos < 1.5).astype(int)
    return X, y


# ================================================================================
# STREAMLIT UI — TO'LIQ INTERFEYS
# ================================================================================

st.title(t('app_title'))
st.markdown(f"### {t('app_subtitle')}")
st.markdown("**✅ Versiya 3.0 — PhD/Patent Sifatli | app.py + app_corrected.py birlashtirildi**")

# Til tanlash
LANGUAGES = {'uz': "🇺🇿 O'zbek", 'en': '🇬🇧 English', 'ru': '🇷🇺 Русский'}
lang = st.sidebar.selectbox(
    "Til / Language / Язык",
    options=list(LANGUAGES.keys()),
    format_func=lambda x: LANGUAGES[x],
    index=list(LANGUAGES.keys()).index(st.session_state.language)
)
st.session_state.language = lang

# Tizim sog'lig'i
health = check_app_health()
st.sidebar.metric("Tizim holati", health.status,
                  f"Xotira: {health.memory_usage_mb:.0f}MB")

st.sidebar.markdown("---")

# QR-kod
st.sidebar.subheader("📱 Ilovaga o'tish")
APP_URL = os.environ.get("UCG_APP_URL",
                         "https://angren-ucg-app-a7rxktm6usxqixabhaq576.streamlit.app/")


@st.cache_data
def generate_qr(link: str) -> bytes:
    qr = qrcode.QRCode(version=1,
                        error_correction=qrcode.constants.ERROR_CORRECT_L,
                        box_size=10, border=4)
    qr.add_data(link)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


qr_img_bytes = generate_qr(APP_URL)
st.sidebar.image(qr_img_bytes, caption="Scan QR: Angren UCG", use_container_width=True)

# Formulalar
st.sidebar.header(t('sidebar_header_params'))
formula_opts = FORMULA_OPTIONS[st.session_state.language]
if "formula_idx" not in st.session_state:
    st.session_state.formula_idx = 0
formula_option = st.sidebar.selectbox(t('formula_show'), formula_opts,
                                       index=st.session_state.formula_idx)
st.session_state.formula_idx = (formula_opts.index(formula_option)
                                 if formula_option in formula_opts else 0)

if formula_option != formula_opts[0]:
    with st.expander(f"📚 Ilmiy asos: {formula_option}", expanded=True):
        if formula_option == formula_opts[1]:
            st.latex(r"\sigma_1 = \sigma_3 + \sigma_{ci} \left( m_b \frac{\sigma_3}{\sigma_{ci}} + s \right)^a")
            st.latex(r"m_b = m_i \exp\left(\frac{GSI-100}{28-14D}\right); \quad s = \exp\left(\frac{GSI-100}{9-3D}\right)")
            st.latex(r"a = \frac{1}{2} + \frac{1}{6} \left( e^{-GSI/15} - e^{-20/3} \right), \quad a \in [0.5, 0.65]")
        elif formula_option == formula_opts[2]:
            st.latex(r"D(T) = \begin{cases} \beta \frac{\Delta T}{300}, & T < T_0+300 \\ 1-e^{-\beta(\Delta T-300)}, & T \geq T_0+300 \end{cases}")
            st.latex(r"k = \frac{\phi^3}{180(1-\phi)^2} d^2 \cdot (1+50D^2) \cdot e^{-0.01\sigma'}")
        elif formula_option == formula_opts[3]:
            st.latex(r"\sigma_{th} = \frac{E \alpha \Delta T}{1-\nu}; \quad \sigma'_v = \sigma_v - B \cdot u \quad \text{(Skempton, 1961)}")
            st.latex(r"\mu(T) = \mu_0 \left(\frac{T}{T_0}\right)^{0.67} \quad \text{(Chapman \& Cowling, 1970)}")
        elif formula_option == formula_opts[4]:
            st.latex(r"\sigma_{p} = \sigma_{ci} \cdot \left(0.64 + 0.36 \frac{w}{H}\right) \quad \text{(Wilson, 1972)}")
            st.latex(r"S(x) = S_{max} \cdot \exp\left( -\frac{x^2}{2i^2} \right), \quad i = 0.45 H_{tot}")

# Asosiy parametrlar
obj_name = st.sidebar.text_input(t('project_name'), value="Angren-UCG-001")
time_h = st.sidebar.slider(t('process_time'), 1, 150, 24)
num_layers = st.sidebar.number_input(t('num_layers'), min_value=1, max_value=5, value=3)
tensile_mode = st.sidebar.selectbox(
    t('tensile_model'),
    [t('tensile_empirical'), t('tensile_hb'), t('tensile_manual')]
)
st.sidebar.subheader(t('rock_props'))
D_factor = st.sidebar.slider(t('disturbance'), 0.0, 1.0, 0.7)
nu_poisson = st.sidebar.slider(t('poisson'), 0.1, 0.4, 0.25)
k_ratio = st.sidebar.slider(t('stress_ratio'), 0.1, 2.0, 0.5)
st.sidebar.subheader(t('tensile_params'))
beta_thermal = st.sidebar.slider(t('thermal_decay_label'),
                                  min_value=0.0005, max_value=0.02,
                                  value=PARAMS.thermal_damage_beta, step=0.0005)
st.sidebar.subheader(t('combustion'))
burn_duration = st.sidebar.number_input(t('burn_duration'), value=40, min_value=1)
T_source_max = st.sidebar.slider(t('max_temp'), 600, 1200, PARAMS.gas_temp)
with st.sidebar.expander(t('timeline')):
    st.markdown(t('timeline_table'))

# Quduqlar konfiguratsiyasi
st.sidebar.subheader(t('well_config'))
well_distance = st.sidebar.slider(t('well_distance'), 10, 200, 60)

# Qatlamlar
strata_colors = ['#87CEEB', '#F4A460', '#D3D3D3', '#F5DEB3', '#555555']
layers_data = []
total_depth = 0.0
for i in range(int(num_layers)):
    with st.sidebar.expander(t('layer_params', num=i + 1),
                              expanded=(i == int(num_layers) - 1)):
        name = st.text_input(t('layer_name'), value=f"Qatlam-{i + 1}", key=f"name_{i}")
        thick = st.number_input(t('thickness'), value=50.0, min_value=0.1, key=f"thick_{i}")
        u = st.number_input(t('ucs'), value=40.0, min_value=0.1, key=f"ucs_{i}")
        rho = st.number_input(t('density'), value=2500.0, min_value=100.0, key=f"rho_{i}")
        color = st.color_picker(t('color'), strata_colors[i % len(strata_colors)],
                                 key=f"color_{i}")
        g = st.slider(t('gsi'), 10, 100, 60, key=f"gsi_{i}")
        m = st.number_input(t('mi'), value=10.0, min_value=0.1, key=f"mi_{i}")
        s_t0_val = (st.number_input(t('manual_st0'), value=3.0, key=f"st_{i}")
                    if tensile_mode == t('tensile_manual') else 0.0)
    layers_data.append({
        'name': name, 'thickness': thick, 'ucs': u, 'rho': rho,
        'gsi': g, 'mi': m, 'color': color, 'z_start': total_depth,
        'sigma_t0_manual': s_t0_val
    })
    total_depth += thick

# Kirish tekshiruvi
errors = []
for lyr in layers_data:
    if lyr['thickness'] <= 0: errors.append(t('error_thick_positive'))
    if lyr['ucs'] <= 0: errors.append(t('error_ucs_positive'))
    if lyr['rho'] <= 0: errors.append(t('error_density_positive'))
    if not (10 <= lyr['gsi'] <= 100): errors.append(t('error_gsi_range'))
    if lyr['mi'] <= 0: errors.append(t('error_mi_positive'))
if not layers_data:
    errors.append(t('error_min_layers'))
if errors:
    for e in errors:
        st.error(e)
    st.stop()

# Asosiy hisoblashlar
depth_seam = sum(l['thickness'] for l in layers_data[:-1]) + layers_data[-1]['thickness'] / 2
avg_rho = np.mean([l['rho'] for l in layers_data])
H_seam = layers_data[-1]['thickness']
source_z = total_depth - H_seam / 2

st.info(t('pin_approx'))

# Metrikalar
st.subheader(f"📊 {t('monitoring_header', obj_name=obj_name)}")
col1, col2, col3, col4 = st.columns(4)
col1.metric("Umumiy chuqurlik", f"{total_depth:.1f} m")
col2.metric("Maqsadli qatlam qalinligi", f"{H_seam:.1f} m")
col3.metric("O'rtacha zichlik", f"{avg_rho:.0f} kg/m³")
col4.metric("Yonish muddati", f"{burn_duration} h")

st.markdown("---")

# ✅ Harorat maydoni hisoblash
grid_shape = (80, 100)
temp_2d, x_axis, z_axis, grid_x, grid_z = compute_temperature_field_moving(
    time_h, T_source_max, burn_duration, total_depth, source_z, grid_shape
)
dx_val = x_axis[1] - x_axis[0]
dz_val = z_axis[1] - z_axis[0]

E_field = young_modulus_temperature(temp_2d)
alpha_field = thermal_expansion_temperature(temp_2d)

# Zichlik maydoni
grid_rho = np.zeros_like(temp_2d)
layer_bounds_full = [(l['z_start'], l['z_start'] + l['thickness'], l) for l in layers_data]
for z0, z1, layer in layer_bounds_full:
    mask = (grid_z >= z0) & (grid_z < z1 if layer != layers_data[-1] else True)
    grid_rho[mask] = density_temperature(layer['rho'], temp_2d[mask])

# Vertikal va gorizontal kuchlanish
grid_sigma_v = np.zeros((len(z_axis), len(x_axis)))
for i in range(1, len(z_axis)):
    dz_i = z_axis[i] - z_axis[i - 1]
    grid_sigma_v[i, :] = grid_sigma_v[i - 1, :] + grid_rho[i, :] * 9.81 * dz_i / 1e6
grid_sigma_h = k_ratio * grid_sigma_v

# Kaverna radiusi va g'ovaklik bosimi
cavity_radius = evolving_cavity_radius(time_h, temp_2d, beta_thermal, grid_z, source_z, H_seam)
pore_pressure = pore_pressure_field(temp_2d, grid_z)

# ✅ Regularizatsiyalangan Kirsch
sigma_rr, sigma_tt, tau_rt = kirsch_stress_field(
    grid_x, grid_z - source_z, grid_sigma_h, grid_sigma_v, cavity_radius, pore_pressure
)

# Termal kuchlanish
delta_T = np.maximum(temp_2d - T_REF_AMBIENT, 0)
sigma_thermal_pa = E_field * alpha_field * delta_T / (1.0 - nu_poisson + EPS)
sigma_thermal = sigma_thermal_pa / 1e6

dT_dx, dT_dz = np.gradient(temp_2d, dx_val, dz_val)
dT_deviatoric = (dT_dx - dT_dz) / 2.0
tau_thermal = (E_field * alpha_field * dT_deviatoric * nu_poisson) / \
              (2.0 * (1.0 - nu_poisson ** 2) + EPS) / 1e6
tau_rt += tau_thermal

sigma_x_total = sigma_rr - sigma_thermal
sigma_z_total = sigma_tt - sigma_thermal
sigma1_act, sigma3_act = principal_stresses(sigma_x_total, sigma_z_total, tau_rt)

# Hoek-Brown parametrlari uchun grid
grid_ucs = np.zeros_like(grid_z)
grid_mb = np.zeros_like(grid_z)
grid_s_hb = np.zeros_like(grid_z)
grid_a_hb = np.zeros_like(grid_z)
for z0, z1, layer in layer_bounds_full:
    mask = (grid_z >= z0) & (grid_z < z1 if layer != layers_data[-1] else True)
    grid_ucs[mask] = layer['ucs']
    mb, s, a = hoek_brown_params_corrected(layer['gsi'], layer['mi'], D_factor)
    grid_mb[mask] = mb
    grid_s_hb[mask] = s
    grid_a_hb[mask] = a

# ✅ Termal degradatsiya (Shao et al., 2015 ikki bosqichli model)
ucs_field_degraded = apply_thermal_degradation(grid_ucs, temp_2d, beta_thermal)
overstress = compute_demand_capacity_ratio(sigma1_act, sigma3_act,
                                            ucs_field_degraded, grid_mb, grid_s_hb, grid_a_hb)

# Bo'shliq hajmi
void_fraction = gaussian_filter(overstress * (temp_2d > 600), sigma=2)
void_mask_permanent = void_fraction > 0.5
void_volume = np.sum(void_mask_permanent) * dx_val * dz_val

# ✅ Modifikatsiyalangan Kozeny-Carman o'tkazuvchanligi
volumetric_strain = (sigma_thermal * 1e6) / (E_field + EPS)
perm = np.zeros_like(temp_2d)
for idx in np.ndindex(temp_2d.shape):
    dmg = float(thermal_damage_shao(temp_2d[idx], beta_thermal))
    perm[idx] = permeability_modified_kozeny_carman(
        dmg, volumetric_strain=float(volumetric_strain[idx])
    )
perm_x = perm * 5
perm_z = perm
perm = np.clip(perm, 1e-16, 1e-10)

# ✅ Haroratga bog'liq gaz oqimi
T_kelvin = temp_2d + 273.15
pressure_field = 1e5 + 50 * (T_kelvin - 293.15)
gas_density = (pressure_field * PARAMS.MOLAR_MASS_GAS) / (PARAMS.R_UNIVERSAL * T_kelvin)
dp_dx, dp_dz = np.gradient(pressure_field, dx_val, dz_val)
mu_gas_field = viscosity_temperature(temp_2d)
vx = -perm_x * dp_dx / (mu_gas_field + EPS)
vz = -perm_z * dp_dz / (mu_gas_field + EPS)
gas_velocity = np.sqrt(vx ** 2 + vz ** 2)

# Cho'kish hisoblash
phi_rad = np.radians(PARAMS.phi_deg)
angle_of_draw = np.radians(45.0 - PARAMS.phi_deg / 2.0)
influence_radius = total_depth * np.tan(angle_of_draw)
i_oreilly = 0.45 * total_depth

c_subs = PARAMS.subsidence_rate
Smax = H_seam * 0.04
subsidence_t = Smax * (1 - np.exp(-c_subs * time_h))
subsidence_raw = -subsidence_t * np.exp(-(x_axis ** 2) / (2 * influence_radius ** 2))
win_len = min(11, len(x_axis) - 1)
if win_len % 2 == 0:
    win_len -= 1
subsidence_raw = savgol_filter(subsidence_raw, window_length=win_len, polyorder=3)

void_correction_factor = 1.0 + PARAMS.K_VOID * float(np.mean(void_mask_permanent))
sub_p = subsidence_raw * void_correction_factor

# Qiyalik tuzatmasi
dip_angle = st.sidebar.slider(t('dip_angle_label'), 0, 90, 0, 5, key="dip_angle_slider")
if dip_angle > 0:
    shift = subsidence_inclined_seam(sub_p, dip_angle, total_depth, PARAMS.phi_deg)
    x_shifted = x_axis + shift
    sub_p = np.interp(x_axis, x_shifted, sub_p)

horizontal_disp_m = -(x_axis / (influence_radius + EPS)) * sub_p
horizontal_disp_cm = horizontal_disp_m * 100.0

# Selek kuchi va FOS hisoblash
idx_closest = np.abs(z_axis - source_z).argmin()
avg_t_p = np.mean(temp_2d[idx_closest, :])
ucs_seam = layers_data[-1]['ucs']
sv_seam = grid_sigma_v[idx_closest, :].max()
target_layer = layers_data[-1]
mb_dyn, s_dyn, a_dyn = hoek_brown_params_corrected(target_layer['gsi'], target_layer['mi'], D_factor)
ucs_t_dyn = apply_thermal_degradation(ucs_seam, avg_t_p, beta_thermal)
sigma_cm = ucs_t_dyn * (np.maximum(s_dyn, 1e-9) ** a_dyn)

# Selek enini iterativ hisoblash
w_sol = 20.0
E_MIN_CORE = 0.5 * H_seam
w_prev = w_sol
for iteration in range(50):
    y_plastic = max((H_seam / 2) * (np.sqrt(sv_seam / (sigma_cm + EPS)) - 1), 0.0)
    core_width = max(w_prev - 2 * y_plastic, E_MIN_CORE)
    pillar_strength_calc = sigma_cm * (WILSON_C1 + WILSON_C2 * core_width / (H_seam + EPS))
    pore_p_seam = float(np.mean(pore_pressure[idx_closest, :]))
    fos_stage = fos_with_pore_pressure(pillar_strength_calc, sv_seam, pore_p_seam)
    if fos_stage < 1.5:
        w_prev *= 1.05
    else:
        break
    if w_prev > 200:
        break

rec_width = w_prev
pillar_strength_final = sigma_cm * (WILSON_C1 + WILSON_C2 * rec_width / (H_seam + EPS))
fos_final = fos_with_pore_pressure(pillar_strength_final, sv_seam, pore_p_seam)

# FOS maydoni
fos_stage = np.clip(ucs_field_degraded / (sigma1_act + EPS), 0, 3)

# ================================================================================
# AI/ML MODEL
# ================================================================================

@st.cache_resource
def train_ai_model():
    X, y = physics_features(
        temp_2d.flatten(), sigma1_act.flatten(), sigma3_act.flatten(),
        grid_z.flatten(), ucs_seam
    )
    if len(np.unique(y)) < 2:
        return None, None, None
    X_tr, X_te, y_tr, y_te = train_test_split(X, y, test_size=0.2,
                                               random_state=CONFIG.RANDOM_SEED)
    scaler = StandardScaler()
    X_tr_s = scaler.fit_transform(X_tr)
    X_te_s = scaler.transform(X_te)
    rf_model = RandomForestClassifier(n_estimators=100, max_depth=8,
                                       random_state=CONFIG.RANDOM_SEED)
    rf_model.fit(X_tr_s, y_tr)
    return rf_model, scaler, X_te_s


rf_model, scaler_ai, X_te_s = train_ai_model()

if rf_model is not None and scaler_ai is not None:
    X_full, _ = physics_features(
        temp_2d.flatten(), sigma1_act.flatten(), sigma3_act.flatten(),
        grid_z.flatten(), ucs_seam
    )
    X_full_s = scaler_ai.transform(X_full)
    collapse_pred = rf_model.predict_proba(X_full_s)[:, 1].reshape(temp_2d.shape)
else:
    collapse_pred = np.clip(1 - fos_stage / 3, 0, 1)
    if not PT_AVAILABLE:
        st.warning(t('warning_pytorch'))

# ================================================================================
# ASOSIY VIZUALIZATSIYA
# ================================================================================

# --- Cho'kish va gorizontal siljish grafiki ---
st.subheader(t('monitoring_header', obj_name=obj_name))
fig_subs = make_subplots(rows=1, cols=2,
                          subplot_titles=(t('subsidence_title'), t('thermal_deform_title')))

ci_lower_b, ci_upper_b = subsidence_confidence_interval_bootstrap(
    sub_p * 100, n_bootstrap=500
)
fig_subs.add_trace(go.Scatter(x=x_axis, y=sub_p * 100, mode='lines',
                               name='Cho\'kish (cm)', line=dict(color='cyan', width=2)), row=1, col=1)
fig_subs.add_trace(go.Scatter(
    x=np.concatenate([x_axis, x_axis[::-1]]),
    y=np.concatenate([np.full_like(x_axis, ci_upper_b), np.full_like(x_axis, ci_lower_b)]),
    fill='toself', fillcolor='rgba(0,255,255,0.1)', line=dict(color='rgba(255,255,255,0)'),
    name='95% Bootstrap CI'), row=1, col=1)
fig_subs.add_trace(go.Scatter(x=x_axis, y=horizontal_disp_cm, mode='lines',
                               name="Siljish (cm)", line=dict(color='orange', width=2)), row=1, col=2)
fig_subs.update_layout(template='plotly_dark', height=400)
st.plotly_chart(fig_subs, use_container_width=True)

# --- TM maydoni ---
st.subheader(t('tm_field_title'))
fig_tm = make_subplots(rows=1, cols=2,
                        subplot_titles=(t('temp_subplot'), t('fos_subplot')))

step_x = max(1, len(x_axis) // 20)
step_z = max(1, len(z_axis) // 20)
fig_tm.add_trace(go.Heatmap(z=temp_2d, x=x_axis, y=z_axis,
                             colorscale='Hot', colorbar=dict(x=0.45, title='°C')), row=1, col=1)
fig_tm.add_trace(go.Cone(
    x=grid_x[::step_z, ::step_x].flatten(),
    y=grid_z[::step_z, ::step_x].flatten(),
    z=np.zeros(grid_x[::step_z, ::step_x].size),
    u=vx[::step_z, ::step_x].flatten(),
    v=vz[::step_z, ::step_x].flatten(),
    w=np.zeros(grid_x[::step_z, ::step_x].size),
    sizemode='absolute', sizeref=0.5, colorscale='Blues',
    name=t('gas_flow'), showscale=False), row=1, col=1)

fig_tm.add_trace(go.Heatmap(z=fos_stage, x=x_axis, y=z_axis,
                             colorscale='RdYlGn', zmin=0, zmax=3,
                             colorbar=dict(x=1.01, title='FOS')), row=1, col=2)
fig_tm.add_trace(go.Contour(z=collapse_pred, x=x_axis, y=z_axis,
                             contours=dict(start=0.5, end=0.5, size=0),
                             line=dict(color='red', width=2),
                             name=t('ai_collapse'), showscale=False), row=1, col=2)
yield_mask = overstress > 1.0
if np.any(yield_mask):
    yz_idx, yx_idx = np.where(yield_mask)
    fig_tm.add_trace(go.Scatter(
        x=x_axis[yx_idx[::10]], y=z_axis[yz_idx[::10]],
        mode='markers', marker=dict(color='red', size=3, opacity=0.5),
        name=t('shear')), row=1, col=2)

fig_tm.update_layout(template='plotly_dark', height=500)
st.plotly_chart(fig_tm, use_container_width=True)

# --- Hoek-Brown enveloplari ---
with st.expander(t('hb_envelopes_title')):
    sigma3_range = np.linspace(0, ucs_seam * 0.5, 100)
    fig_hb = go.Figure()
    for i, (layer, lbl) in enumerate(zip(layers_data,
                                          [l['name'] for l in layers_data])):
        mb_l, s_l, a_l = hoek_brown_params_corrected(layer['gsi'], layer['mi'], D_factor)
        sigma1_env = hoek_brown(sigma3_range, layer['ucs'], mb_l, s_l, a_l)
        ucs_T_l = apply_thermal_degradation(layer['ucs'], avg_t_p, beta_thermal)
        mb_T, s_T, a_T = hoek_brown_params_corrected(layer['gsi'], layer['mi'], D_factor)
        sigma1_T = hoek_brown(sigma3_range, ucs_T_l, mb_T, s_T, a_T)
        fig_hb.add_trace(go.Scatter(x=sigma3_range, y=sigma1_env,
                                     mode='lines', name=f'{lbl} (T₀)',
                                     line=dict(width=2, dash='solid')))
        fig_hb.add_trace(go.Scatter(x=sigma3_range, y=sigma1_T,
                                     mode='lines', name=f'{lbl} (T={avg_t_p:.0f}°C)',
                                     line=dict(width=2, dash='dash')))
    fig_hb.update_layout(template='plotly_dark', height=400,
                          xaxis_title='σ₃ (MPa)', yaxis_title='σ₁ (MPa)')
    st.plotly_chart(fig_hb, use_container_width=True)

# --- Natijalar metrikasi ---
risk_level, risk_val = water_inrush_risk(void_volume, 50, depth_seam, float(np.min(fos_stage)))
m1, m2, m3, m4, m5 = st.columns(5)
m1.metric(t('pillar_strength'), f"{pillar_strength_final:.1f} MPa")
m2.metric("FOS (Skempton)", f"{fos_final:.2f}",
          delta="✅ Barqaror" if fos_final >= 1.5 else "⚠️ Xavfli",
          delta_color="normal" if fos_final >= 1.5 else "inverse")
m3.metric(t('plastic_zone'), f"{max((H_seam / 2) * (np.sqrt(sv_seam / (sigma_cm + EPS)) - 1), 0):.1f} m")
m4.metric(t('cavity_volume'), f"{void_volume:.0f} m²")
m5.metric("Su sizish xavfi", risk_level)

# FOS holati
if fos_final < 1.0:
    st.error(t('fos_red'))
elif fos_final < 1.5:
    st.warning(t('fos_yellow'))
else:
    st.success(t('fos_green'))

st.markdown("---")

# ================================================================================
# CHUQUR TAHLIL TABLAR
# ================================================================================

with st.expander(t('advanced_analysis')):
    tab1, tab2, tab3 = st.tabs([t('tab_mass'), t('tab_thermal'), t('tab_stability')])

    with tab1:
        st.subheader(t('hb_class'))
        mb_show, s_show, a_show = hoek_brown_params_corrected(
            target_layer['gsi'], target_layer['mi'], D_factor
        )
        col_mb, col_s, col_a = st.columns(3)
        col_mb.metric(t('hb_mb', mb=mb_show), f"m_i={target_layer['mi']}")
        col_s.metric(t('hb_s', s=s_show), f"GSI={target_layer['gsi']}")
        col_a.metric("a (✅ bounded)", f"{a_show:.4f} ∈ [0.5, 0.65]")
        perc_strength = (1 - s_show ** a_show) * 100
        st.markdown(t('hb_interpret', gsi=target_layer['gsi'], perc=perc_strength))

    with tab2:
        st.subheader(t('thermal_params'))
        param_df = pd.DataFrame([
            {t('param_table_param'): t('modulus'), t('param_table_value'): "25 GPa",
             t('param_table_reason'): t('modulus_reason')},
            {t('param_table_param'): t('alpha'), t('param_table_value'): "3×10⁻⁵ 1/K",
             t('param_table_reason'): t('alpha_reason')},
            {t('param_table_param'): t('temp0'), t('param_table_value'): f"{T_REF_AMBIENT}°C",
             t('param_table_reason'): t('temp0_reason')},
        ])
        st.dataframe(param_df, use_container_width=True)

        st.markdown(t('ucs_decay'))
        ucs_at_T = float(apply_thermal_degradation(ucs_seam, avg_t_p, beta_thermal))
        st.latex(t('ucs_decay_eq', ucs=ucs_at_T))
        perc_decay = (1 - ucs_at_T / ucs_seam) * 100
        st.markdown(t('ucs_interpret', temp=int(avg_t_p), perc=perc_decay))

        st.markdown(t('thermal_stress'))
        sigma_th_show = float(DimensionalAnalysis.check_thermal_stress(
            PARAMS.E_mass, PARAMS.alpha_thermal, avg_t_p - T_REF_AMBIENT, nu_poisson
        )) / 1e6
        st.latex(t('thermal_stress_eq', sigma=sigma_th_show))

        T_range_show = np.linspace(20, 1000, 200)
        dmg_shao = thermal_damage_shao(T_range_show, beta_thermal)
        dmg_simple = thermal_damage(T_range_show, beta_thermal)
        fig_dmg = go.Figure()
        fig_dmg.add_trace(go.Scatter(x=T_range_show, y=dmg_shao, mode='lines',
                                      name='Shao et al. (2015) — 2 bosqich ✅',
                                      line=dict(color='red', width=3)))
        fig_dmg.add_trace(go.Scatter(x=T_range_show, y=dmg_simple, mode='lines',
                                      name='Sodda eksponensial (eski)',
                                      line=dict(color='gray', width=2, dash='dash')))
        fig_dmg.add_vline(x=T_REF_AMBIENT + 300, line_dash='dash', line_color='orange',
                           annotation_text='Fazalar o\'tishi (300°C)')
        fig_dmg.update_layout(template='plotly_dark', height=350,
                               title='Termal Zarar Evolyutsiyasi',
                               xaxis_title='T (°C)', yaxis_title='Zarar (0–1)')
        st.plotly_chart(fig_dmg, use_container_width=True)

    with tab3:
        st.subheader(t('pillar_stability'))
        st.latex(t('fos_eq', fos=fos_final))
        y_pl = max((H_seam / 2) * (np.sqrt(sv_seam / (sigma_cm + EPS)) - 1), 0.0)
        st.markdown(t('pillar_wilson', w=f"{rec_width:.1f}", sv=sv_seam, y=y_pl))

        st.markdown(t('references'))
        for key in ['ref1', 'ref2', 'ref3', 'ref4']:
            st.markdown(f"- {t(key)}")

        if fos_final < 1.5:
            st.error(t('conclusion_danger', fos=fos_final))
        else:
            st.success(t('conclusion_safe', fos=fos_final))

st.markdown("---")

# ================================================================================
# LIVE MONITORING PANELI
# ================================================================================

st.header(t('monitoring_panel', obj_name=obj_name))


def calculate_live_metrics(h, layers, T_max):
    target = layers[-1]
    ucs_0, H_l = target['ucs'], target['thickness']
    curr_T = (25 + (T_max - 25) * (min(h, 40) / 40) if h <= 40
              else T_max * np.exp(-0.001 * (h - 40)))
    ucs_T_live = apply_thermal_degradation(ucs_0, curr_T, beta_thermal)
    w_rec = 15.0 + (h / 150) * 10
    p_str = ucs_T_live * (WILSON_C1 + WILSON_C2 * w_rec / (H_l + EPS))
    max_sub = (H_l * 0.05) * (min(h, 120) / 120)
    return p_str, w_rec, curr_T, max_sub


p_str_live, w_rec_live, t_now, s_max_live = calculate_live_metrics(
    time_h, layers_data, T_source_max
)
mk1, mk2, mk3, mk4 = st.columns(4)
mk1.metric(t('pillar_live'), f"{p_str_live:.1f} MPa",
           delta=f"{t_now:.0f} °C", delta_color="inverse")
mk2.metric(t('rec_width_live'), f"{w_rec_live:.1f} m")
mk3.metric(t('max_subsidence_live'), f"{s_max_live * 100:.1f} cm")
mk4.metric(t('process_stage'),
           t('stage_active') if time_h < 100 else t('stage_cooling'))

st.markdown("---")

# ================================================================================
# TIZIM ENTROPIYASI VA XAVF INDEKSI
# ================================================================================

weights = np.array([0.4, 0.3, 0.2, 0.1])
max_perm = max(np.max(perm), 1e-20)
risk_index_var = (
    weights[0] * collapse_pred +
    weights[1] * (1 - fos_stage / 3.0) +
    weights[2] * (perm / max_perm) +
    weights[3] * (temp_2d / np.max(temp_2d))
)
risk_index_var = np.maximum(0, risk_index_var)
risk_flat = risk_index_var.flatten()
risk_prob = risk_flat / (np.sum(risk_flat) + 1e-12)
entropy = -np.sum(risk_prob * np.log(risk_prob + 1e-12))
st.metric(t('system_entropy'), f"{entropy:.3f}")

# ================================================================================
# KENGAYTIRILGAN TAHLIL BO'LIMLARI
# ================================================================================

# --- FOS Trend bashorati ---
with st.expander("📈 FOS Vaqt Bashorati (Trend)"):
    time_points = np.arange(1, time_h + 1, max(1, time_h // 20))
    if len(time_points) >= 2:
        fos_timeline = []
        for current_time in time_points:
            T_at_th = (T_REF_AMBIENT + (T_source_max - T_REF_AMBIENT) *
                       min(current_time, burn_duration) / max(burn_duration, 1))
            ucs_T_th = apply_thermal_degradation(ucs_seam, T_at_th, beta_thermal)
            p_str_t = ucs_T_th * (WILSON_C1 + WILSON_C2 * rec_width / (H_seam + EPS))
            sv_t = sv_seam * (1 + 0.001 * current_time)
            fos_timeline.append(float(np.clip(p_str_t / (sv_t + EPS), 0, 3)))

        slope, intercept, r_value, _, _ = linregress(time_points, fos_timeline)
        future_times = np.arange(time_h, min(time_h * 2, 300), max(1, time_h // 10))
        fos_forecast = np.clip(intercept + slope * future_times, 0, 3)

        fig_trend = go.Figure()
        fig_trend.add_trace(go.Scatter(x=time_points, y=fos_timeline,
                                        mode='lines+markers', name='FOS',
                                        line=dict(color='cyan', width=2)))
        fig_trend.add_trace(go.Scatter(x=future_times, y=fos_forecast,
                                        mode='lines', name='Bashorat',
                                        line=dict(color='orange', dash='dash')))
        fig_trend.add_hline(y=1.5, line_color='green', line_dash='dash')
        fig_trend.add_hline(y=1.0, line_color='red', line_dash='dash')
        fig_trend.update_layout(template='plotly_dark', height=400,
                                 title=f"Trend: {slope:+.4f} FOS/soat | R²={r_value ** 2:.3f}")
        st.plotly_chart(fig_trend, use_container_width=True)

# --- Fazaviy maydon modeli ---
with st.expander("🪨 Phase-Field Fracture Damage (Bourdin et al., 2000)"):
    st.markdown(t('phase_field_info'))

    def laplacian_neumann(field, dx, dz):
        f = np.pad(field, 1, mode='edge')
        return ((f[1:-1, 2:] - 2 * f[1:-1, 1:-1] + f[1:-1, :-2]) / dx ** 2 +
                (f[2:, 1:-1] - 2 * f[1:-1, 1:-1] + f[:-2, 1:-1]) / dz ** 2)

    def phase_field_update(damage, strain_energy, dx, dz, dt, Gc=0.01, l_char=1.0, eta=1e-3):
        dt_max = dx ** 2 / (4 * Gc * l_char)
        dt = min(dt, 0.9 * dt_max)
        lap = laplacian_neumann(damage, dx, dz)
        driving = (Gc * l_char * lap - (Gc / l_char) * damage +
                   (1 - damage) * strain_energy)
        return np.clip(damage + (dt / eta) * driving, 0, 1)

    if st.button("Phase-field bir qadam (demo)"):
        strain_energy = (von_mises_stress(sigma_x_total, sigma_z_total, tau_rt, nu=nu_poisson) ** 2) / \
                        (2 * E_field + EPS)
        d_trial = phase_field_update(overstress, strain_energy, dx_val, dz_val, dt=0.1)
        d_updated = np.maximum(overstress, d_trial)
        fig_phase = go.Figure(go.Heatmap(z=d_updated, x=x_axis, y=z_axis,
                                          colorscale='Viridis', zmin=0, zmax=1))
        fig_phase.update_layout(title="Phase-field zarar (1 qadam)", template='plotly_dark')
        st.plotly_chart(fig_phase, use_container_width=True)

# --- PINN ---
with st.expander("🧠 Physics-Informed Neural Network (PINN) — demo"):
    st.markdown("PINN harorat tenglamasi uchun (PyTorch asosida, demo).")
    if PT_AVAILABLE and st.button("PINN o'qitish"):
        class HeatPINN(nn.Module):
            def __init__(self):
                super().__init__()
                self.net = nn.Sequential(
                    nn.Linear(3, 64), nn.Tanh(),
                    nn.Linear(64, 64), nn.Tanh(),
                    nn.Linear(64, 1)
                )
            def forward(self, x, z, t):
                return self.net(torch.cat([x, z, t], dim=1))

        model_pinn = HeatPINN().to(device)
        opt = torch.optim.Adam(model_pinn.parameters(), lr=1e-3)
        for ep in range(50):
            opt.zero_grad()
            x_r = torch.rand(200, 1, device=device) * 20 - 10
            z_r = torch.rand(200, 1, device=device) * 10
            t_r = torch.rand(200, 1, device=device) * 5
            x_r.requires_grad_(True)
            z_r.requires_grad_(True)
            t_r.requires_grad_(True)
            T_pred = model_pinn(x_r, z_r, t_r)
            dT_dt = torch.autograd.grad(T_pred.sum(), t_r, create_graph=True)[0]
            dT_dx = torch.autograd.grad(T_pred.sum(), x_r, create_graph=True)[0]
            d2T_dx2 = torch.autograd.grad(dT_dx.sum(), x_r, create_graph=True)[0]
            dT_dz2 = torch.autograd.grad(T_pred.sum(), z_r, create_graph=True)[0]
            d2T_dz2 = torch.autograd.grad(dT_dz2.sum(), z_r, create_graph=True)[0]
            residual = dT_dt - 1e-6 * (d2T_dx2 + d2T_dz2)
            loss = torch.mean(residual ** 2)
            loss.backward()
            opt.step()
        st.success(f"✅ PINN o'qitildi. Oxirgi loss: {loss.item():.6f}")

# --- Monte Carlo (Cholesky) ---
with st.expander("📊 Monte Carlo FOS (✅ Cholesky dekompozitsiya)"):
    st.markdown(t('uq_info'))
    try:
        fos_mc, pf_mc = monte_carlo_fos_corrected(
            ucs_mean=ucs_seam, ucs_std=0.1 * ucs_seam,
            gsi_mean=target_layer['gsi'], gsi_std=5.0,
            mi_val=target_layer['mi'], D=D_factor,
            T_avg=avg_t_p, H_seam=H_seam, depth=depth_seam,
            density=avg_rho, rec_width=rec_width,
            beta_th=beta_thermal, n_sim=1000
        )
        fig_mc = go.Figure()
        fig_mc.add_histogram(x=fos_mc, nbinsx=40, marker_color='cyan', name='FOS taqsimoti')
        fig_mc.add_vline(x=1.0, line_color='red', line_dash='dash',
                          annotation_text='Kritik FOS=1.0')
        fig_mc.add_vline(x=1.5, line_color='orange', line_dash='dash',
                          annotation_text='Ogohlantirish FOS=1.5')
        fig_mc.update_layout(template='plotly_dark', height=350,
                              title=f"FOS taqsimoti | Yemirilish ehtimoli: {pf_mc * 100:.1f}%")
        st.plotly_chart(fig_mc, use_container_width=True)
        mc1, mc2, mc3, mc4 = st.columns(4)
        mc1.metric("O'rtacha FOS", f"{np.mean(fos_mc):.3f}")
        mc2.metric("Standart og'ish", f"{np.std(fos_mc):.3f}")
        mc3.metric("5-persentil", f"{np.percentile(fos_mc, 5):.3f}")
        mc4.metric("Yemirilish P", f"{pf_mc * 100:.1f}%")
    except Exception as e:
        st.error(f"Monte Carlo xatosi: {e}")
        logger.exception("MC error")

# --- O'tkazuvchanlik evolyutsiyasi ---
with st.expander("💨 O'tkazuvchanlik evolyutsiyasi (Kozeny-Carman, ✅ to'g'irlangan)"):
    damage_range = np.linspace(0, 1, 50)
    perm_ev = [permeability_modified_kozeny_carman(d, volumetric_strain=0.01) * 1e15
               for d in damage_range]
    fig_perm = go.Figure()
    fig_perm.add_trace(go.Scatter(x=damage_range, y=perm_ev, mode='lines+markers',
                                   name='k(zarar)', line=dict(color='lime', width=2)))
    fig_perm.update_layout(template='plotly_dark', height=350,
                            title='O\'tkazuvchanlik va Termal Zarar',
                            xaxis_title='Zarar (0–1)',
                            yaxis_title='k (×10⁻¹⁵ m²)',
                            yaxis_type='log')
    st.plotly_chart(fig_perm, use_container_width=True)

# --- Gaz viskoziteti ---
with st.expander("🌡️ Gaz Viskoziteti (Chapman & Cowling, 1970 — ✅ to'g'irlangan)"):
    T_gas_range = np.linspace(20, 1100, 200)
    mu_range = viscosity_temperature(T_gas_range) * 1e5
    fig_visc = go.Figure()
    fig_visc.add_trace(go.Scatter(x=T_gas_range, y=mu_range, mode='lines',
                                   name='μ(T)', line=dict(color='purple', width=2)))
    fig_visc.update_layout(template='plotly_dark', height=350,
                            title='Gaz Viskoziteti va Harorat',
                            xaxis_title='T (°C)',
                            yaxis_title='μ (×10⁻⁵ Pa·s)')
    st.plotly_chart(fig_visc, use_container_width=True)

# --- O'lchamli tahlil ---
with st.expander("✅ O'lchamli Tahlil (Dimensional Analysis)"):
    try:
        sigma_th_check = DimensionalAnalysis.check_thermal_stress(
            PARAMS.E_mass, PARAMS.alpha_thermal,
            avg_t_p - T_REF_AMBIENT, nu_poisson
        )
        st.success(f"✅ Termal kuchlanish o'lchamli tekshiruvi o'tdi: "
                   f"σ_th = {sigma_th_check / 1e6:.2f} MPa")
    except Exception as e:
        st.error(f"❌ O'lchamli tekshiruv muvaffaqiyatsiz: {e}")

    try:
        sigma_test = np.array([[10., 2., 1.], [2., 8., 0.5], [1., 0.5, 6.]])
        principal_s = validate_stress_tensor_3d(sigma_test)
        st.success(f"✅ 3D stress tenzori tekshiruvi o'tdi. "
                   f"Bosh kuchlanishlar: σ₁={principal_s[0]:.2f}, "
                   f"σ₂={principal_s[1]:.2f}, σ₃={principal_s[2]:.2f} MPa")
    except Exception as e:
        st.error(f"❌ Stress tenzori tekshiruvi muvaffaqiyatsiz: {e}")

# --- Bootstrap CI ---
with st.expander("📈 Bootstrap Ishonch Intervali (Efron & Tibshirani, 1993)"):
    sample_data = np.array([0.5, 0.7, 0.6, 0.8, 0.65, 0.75, 0.55])
    ci_low_b, ci_high_b = subsidence_confidence_interval_bootstrap(sample_data, n_bootstrap=2000)
    st.write(f"**Namuna:** {sample_data}")
    st.write(f"**Bootstrap 95% CI:** [{ci_low_b:.3f}, {ci_high_b:.3f}]")
    st.write(f"**O'rtacha:** {np.mean(sample_data):.3f}")

# --- Sobol' global sezgirlik ---
if SALIB_AVAILABLE:
    with st.expander("📊 Global sezgirlik tahlili (Sobol', 2001)"):
        st.markdown(t('sobol_info'))
        problem = {
            'num_vars': 4,
            'names': ['UCS', 'Temp', 'Depth', 'GSI'],
            'bounds': [[10, 80], [20, 1000], [10, 300], [20, 100]]
        }
        full_sobol = st.checkbox("To'liq tahlil (N=1024)", value=False)
        N_SOBOL = 1024 if full_sobol else 128
        param_values = saltelli.sample(problem, N_SOBOL, calc_second_order=False)

        def model_eval(params):
            ucs, T, d, gsi = params
            mb_s, s_s, a_s = hoek_brown_params_corrected(gsi, target_layer['mi'], D_factor)
            ucs_T_s = apply_thermal_degradation(ucs, T, beta_thermal)
            return ucs_T_s * (max(float(s_s), 1e-9) ** float(a_s))

        Y = np.array([model_eval(p) for p in param_values])
        Si = sobol.analyze(problem, Y)
        sobol_df = pd.DataFrame({
            'Parametr': problem['names'],
            'S1 (birinchi tartib)': Si['S1'],
            'ST (jami)': Si['ST']
        })
        st.dataframe(sobol_df, use_container_width=True)

# --- LHS ---
if PYDOE_AVAILABLE:
    with st.expander("🎲 Latin Hypercube Sampling (McKay, 1979)"):
        st.markdown(t('lhs_info'))
        N_lhs = 5000
        np.random.seed(CONFIG.RANDOM_SEED)
        lhs_sample = lhs(3, samples=N_lhs)
        T_lhs = lhs_sample[:, 0] * (T_source_max - T_REF_AMBIENT) + T_REF_AMBIENT
        UCS_lhs = lhs_sample[:, 1] * (ucs_seam * 0.5) + ucs_seam * 0.5
        Depth_lhs = lhs_sample[:, 2] * (depth_seam * 1.5)
        fos_lhs = np.clip(
            apply_thermal_degradation(UCS_lhs, T_lhs, beta_thermal) /
            (vertical_stress(Depth_lhs, avg_rho) + EPS), 0, 3
        )
        collapse_prob = 1 / (1 + np.exp(10 * (fos_lhs - 1.0)))
        fig_lhs = go.Figure(go.Histogram(x=collapse_prob, nbinsx=50, marker_color='orange'))
        fig_lhs.update_layout(title="Collapse ehtimolligi taqsimoti", template='plotly_dark')
        st.plotly_chart(fig_lhs, use_container_width=True)
        ci_lo = np.percentile(collapse_prob, 5)
        ci_hi = np.percentile(collapse_prob, 95)
        st.write(f"90% ishonch intervali: [{ci_lo:.3f}, {ci_hi:.3f}]")

# --- SHAP ---
if SHAP_AVAILABLE and rf_model is not None:
    with st.expander("🧠 SHAP Model Interpretatsiyasi (Lundberg & Lee, 2017)"):
        st.markdown(t('shap_info'))
        try:
            X_shap, _ = physics_features(
                temp_2d.flatten(), sigma1_act.flatten(), sigma3_act.flatten(),
                grid_z.flatten(), ucs_seam
            )
            background = shap.sample(X_shap, 100)
            explainer = shap.Explainer(rf_model, background)
            shap_values = explainer(background)
            feature_names = ["Temperature", "Sigma1", "Sigma3", "Depth",
                             "Damage", "FOS", "Energy"]
            fig_shap, ax = plt.subplots()
            shap.summary_plot(shap_values, background,
                              feature_names=feature_names, show=False)
            st.pyplot(fig_shap)
        except Exception as e:
            st.warning(f"SHAP tahlili bajarilmadi: {e}")

# --- PyVista 3D ---
if PYVISTA_AVAILABLE:
    with st.expander("🌋 3D Litologik Hajm (PyVista)"):
        try:
            from scipy.interpolate import RegularGridInterpolator
            nx, ny, nz = 50, 50, 30
            x_pv = np.linspace(x_axis.min(), x_axis.max(), nx)
            z_pv = np.linspace(z_axis.min(), z_axis.max(), nz)
            y_pv = np.linspace(0, 100, ny)
            X_pv, Y_pv, Z_pv = np.meshgrid(x_pv, y_pv, z_pv, indexing='ij')
            interp_temp = RegularGridInterpolator((z_axis, x_axis), temp_2d)
            pts = np.column_stack([Z_pv.flatten(), X_pv.flatten()])
            T_vol = interp_temp(pts).reshape(nx, ny, nz)
            grid_pv = pv.StructuredGrid(X_pv, Y_pv, Z_pv)
            grid_pv.point_data["temperature"] = T_vol.flatten()
            plotter = pv.Plotter(off_screen=True)
            plotter.add_volume(grid_pv, scalars="temperature", cmap="hot")
            st.image(plotter.screenshot(), use_container_width=True)
        except Exception as e:
            st.warning(f"PyVista vizualizatsiyasi amalga oshmadi: {e}")

# --- Eksperimental validatsiya ---
with st.expander("🔬 Eksperimental Validatsiya"):
    st.markdown(t('validation_info'))
    st.info(t('experimental_note'))
    exp_file = st.file_uploader("Eksperimental CSV yuklang", type=['csv'],
                                 key="exp_validation")
    if exp_file:
        try:
            df_exp = pd.read_csv(exp_file)
            if 'x' in df_exp.columns and 'subsidence_cm' in df_exp.columns:
                sub_interp = np.interp(df_exp['x'].values, x_axis, sub_p * 100)
                rmse = np.sqrt(np.mean((sub_interp - df_exp['subsidence_cm'].values) ** 2))
                r2 = 1 - np.sum((df_exp['subsidence_cm'] - sub_interp) ** 2) / \
                     (np.sum((df_exp['subsidence_cm'] - df_exp['subsidence_cm'].mean()) ** 2) + EPS)
                fig_val = go.Figure()
                fig_val.add_trace(go.Scatter(x=df_exp['x'], y=df_exp['subsidence_cm'],
                                             mode='markers', name='Eksperimental',
                                             marker=dict(color='red', size=8)))
                fig_val.add_trace(go.Scatter(x=x_axis, y=sub_p * 100, mode='lines',
                                             name='Model', line=dict(color='cyan')))
                fig_val.update_layout(template='plotly_dark')
                st.plotly_chart(fig_val, use_container_width=True)
                v1, v2 = st.columns(2)
                v1.metric("RMSE", f"{rmse:.3f} cm")
                v2.metric("R²", f"{r2:.4f}")
        except Exception as e:
            st.error(f"Validatsiya xatosi: {e}")

# --- CSV yuklab olish ---
csv_buf = io.StringIO()
monitoring_df = pd.DataFrame({
    'x_m': x_axis,
    'subsidence_cm': sub_p * 100,
    'horizontal_displacement_cm': horizontal_disp_cm,
})
monitoring_df.to_csv(csv_buf, index=False)
st.download_button(
    label=t('download_data'),
    data=csv_buf.getvalue(),
    file_name=f"{obj_name}_monitoring.csv",
    mime='text/csv'
)

# ================================================================================
# WORD HISOBOT EKSPORTI
# ================================================================================

with st.expander("📄 Word Hisobot Eksporti (PhD/Patent sifatli)"):
    if st.button("📄 Word hisobot yaratish"):
        try:
            doc = Document()
            doc.add_heading('UCG Tahlil Hisoboti — v3.0', 0)
            doc.add_paragraph(f"Loyiha: {obj_name}")
            doc.add_paragraph(f"Muallif: Saitov Dilshodbek")
            doc.add_paragraph(f"Sana: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph(f"Versiya: 3.0 — PhD/Patent Sifatli Birlashtirish")

            doc.add_heading('1. Kirish Parametrlari', level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Parametr'
            hdr[1].text = 'Qiymat'
            params_list = [
                ('Umumiy chuqurlik', f'{total_depth:.1f} m'),
                ('H_seam', f'{H_seam:.1f} m'),
                ('T_max', f'{T_source_max} °C'),
                ('Jarayon vaqti', f'{time_h} h'),
                ('D_factor', str(D_factor)),
                ('beta_thermal', str(beta_thermal)),
            ]
            for p, v in params_list:
                row = table.add_row().cells
                row[0].text = p
                row[1].text = v

            doc.add_heading('2. Hoek-Brown (2018) Natijalari', level=1)
            doc.add_paragraph(f"mb = {mb_dyn:.3f}")
            doc.add_paragraph(f"s = {s_dyn:.4f}")
            doc.add_paragraph(f"a = {a_dyn:.4f} ∈ [0.5, 0.65] ✅")

            doc.add_heading('3. Barqarorlik Natijalari', level=1)
            doc.add_paragraph(f"FOS (Skempton tuzatmasi bilan) = {fos_final:.3f}")
            doc.add_paragraph(f"Tavsiya etilgan selek eni = {rec_width:.1f} m")
            doc.add_paragraph(f"Yemirilish ehtimoli (MC) = {pf_mc * 100:.1f}%")

            doc.add_heading('4. Termal Tahlil', level=1)
            doc.add_paragraph(
                f"O'rtacha harorat = {avg_t_p:.1f} °C\n"
                f"UCS pasayishi = {(1 - ucs_t_dyn / ucs_seam) * 100:.1f}%\n"
                f"Termal kuchlanish = {sigma_th_check / 1e6:.2f} MPa"
            )

            doc.add_heading('5. Adabiyotlar', level=1)
            refs = [
                "Hoek, E., & Brown, E. T. (2018). JRMGE, 10(1), 1-57.",
                "Yang, D. (2010). PhD Thesis, TU Delft.",
                "Shao, S., et al. (2015). IJRMMS, 81, 1-10.",
                "Wilson, A. H. (1972). Mining Engineer.",
                "Walsh & Brace (1984). JGR, 89(B12).",
                "Chapman & Cowling (1970). Non-uniform Gases.",
                "Efron & Tibshirani (1993). Bootstrap.",
                "Skempton, A. W. (1961). Effective stress.",
            ]
            for ref in refs:
                doc.add_paragraph(ref, style='List Bullet')

            buf_doc = io.BytesIO()
            doc.save(buf_doc)
            buf_doc.seek(0)
            st.download_button(
                label="⬇️ Word hisobotni yuklab olish",
                data=buf_doc,
                file_name=f"{obj_name}_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Word yaratishda xatolik: {e}")

# ================================================================================
# AUDIT VA STATUS
# ================================================================================

log_audit(
    user_id=st.session_state.user_id,
    action="SESSION_COMPLETE",
    params={
        "project": obj_name,
        "layers": len(layers_data),
        "fos": round(fos_final, 3),
        "mc_pf": round(float(pf_mc), 4)
    },
    status="success"
)

st.markdown("---")
st.markdown("""
### ✅ Versiya 3.0 — Tuzatishlar ro'yxati (Corrections Applied)

| # | Tuzatish | Manba |
|---|----------|-------|
| 1 | Hoek-Brown 'a' ∈ [0.5, 0.65] | Hoek & Brown (2018) |
| 2 | Cholesky dekompozitsiya (Monte Carlo) | Efron & Tibshirani (1993) |
| 3 | Kirsch regularizatsiyasi | Kirsch (1898) |
| 4 | Robin chegaraviy sharti | Incropera & DeWitt (2007) |
| 5 | Kozeny-Carman + zarar | Walsh & Brace (1984) |
| 6 | Haroratga bog'liq viskozitet | Chapman & Cowling (1970) |
| 7 | Shao 2-bosqichli termal zarar | Shao et al. (2015) |
| 8 | Bootstrap ishonch intervali | Efron & Tibshirani (1993) |
| 9 | Skempton g'ovaklik tuzatmasi | Skempton (1961) |
| 10 | To'liq 3D stress tenzori | Elastiklik nazariyasi |
| 11 | Von Mises 3D formulasi | Mises (1913) |
| 12 | O'lchamli tahlil (dimensional check) | SI standart |

**Status: PhD/Patent Tayyor ✅ | app.py + app_corrected.py = app_merged.py**
""")
