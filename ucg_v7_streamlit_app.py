    ax.pie([exo/total*100, endo/total*100], labels=['Ekzotermik', 'Endotermik'],
           colors=['#E74C3C', '#3498DB'], autopct='%1.1f%%', startangle=90, textprops={'fontsize': 12})
    ax.set_title('Issiqlik Ulushi', fontsize=14, fontweight='bold')
    plt.tight_layout()
    st.pyplot(fig)

def plot_ucg_zones(df):
    fig, ax = plt.subplots(figsize=(10, 5))
    x, conv = df['time_elapsed'], df['char_conversion']
    ax.fill_between(x, 0, 1, where=conv < 0.3, color='lightcoral', alpha=0.3, label="Boshlang'ich")
    ax.fill_between(x, 0, 1, where=(conv >= 0.3) & (conv < 0.7), color='khaki', alpha=0.3, label='Faol')
    ax.fill_between(x, 0, 1, where=conv >= 0.7, color='lightgreen', alpha=0.3, label="To'liq")
    ax.plot(x, conv, color='darkred', linewidth=2.5)
    ax.set_title('UCG Jarayon Zonalari', fontsize=14, fontweight='bold')
    ax.set_xlabel('Vaqt (s)'); ax.set_ylabel('Konversiya')
    ax.set_ylim(0, 1.05); ax.legend(fontsize=9, loc='center right'); ax.grid(True, linestyle=':')
    plt.tight_layout()
    st.pyplot(fig)

def plot_3d_surface(df):
    fig = plt.figure(figsize=(12, 8))
    ax = fig.add_subplot(111, projection='3d')
    t = df['time_elapsed'].values
    conv = df['char_conversion'].values
    temp = df['temperature'].values
    from scipy.interpolate import griddata
    T_grid, C_grid = np.meshgrid(np.linspace(t.min(), t.max(), 50),
                                   np.linspace(conv.min(), conv.max(), 50))
    Temp_grid = griddata(np.column_stack([t, conv]), temp, (T_grid, C_grid), method='linear')
    surf = ax.plot_surface(T_grid, C_grid, Temp_grid, cmap='hot', alpha=0.85, edgecolor='none')
    ax.set_xlabel('Vaqt (s)'); ax.set_ylabel('Konversiya'); ax.set_zlabel('Harorat (K)')
    ax.set_title('3D: Harorat × Vaqt × Konversiya', fontsize=14, fontweight='bold')
    fig.colorbar(surf, ax=ax, shrink=0.5, aspect=10, label='Harorat (K)')
    plt.tight_layout()
    st.pyplot(fig)

def plot_novelty_index(df):
    fig, ax = plt.subplots(figsize=(10, 5))
    if 'novelty_index' in df.columns:
        ax.plot(df['time_elapsed'], df['novelty_index'], color='#2C3E50', linewidth=2.5)
        ax.fill_between(df['time_elapsed'], df['novelty_index'], alpha=0.15, color='#2C3E50')
    ax.set_title('UCG Novelty Index', fontsize=14, fontweight='bold')
    ax.set_xlabel('Vaqt (s)'); ax.set_ylabel('Novelty (%)')
    ax.grid(True, linestyle=':')
    plt.tight_layout()
    st.pyplot(fig)

def plot_monte_carlo(mc_summary, n_sim):
    fig, axs = plt.subplots(2, 2, figsize=(16, 10))
    fig.suptitle(f'Monte Carlo Noaniqlik Tahlili ({n_sim} simulyatsiya)', fontsize=16, fontweight='bold')
    x_mc = np.arange(mc_summary['temperature']['mean'].shape[0])
    items = [(axs[0,0], 'temperature', 'Harorat (K)', 'red'),
             (axs[0,1], 'char_conversion', 'Konversiya', 'darkred'),
             (axs[1,0], 'porosity', 'Porozlik', 'blue'),
             (axs[1,1], 'novelty_index', 'Novelty (%)', '#2C3E50')]
    for ax, key, ylabel, color in items:
        s = mc_summary[key]
        ax.plot(x_mc, s['mean'], color=color, linewidth=2, label="O'rtacha")
        ax.fill_between(x_mc, s['ci95_low'], s['ci95_high'], alpha=0.25, color=color, label='95% CI')
        ax.set_title(f'MC: {ylabel}', fontsize=12, fontweight='bold')
        ax.set_ylabel(ylabel); ax.legend(fontsize=9); ax.grid(True, linestyle=':')
    plt.tight_layout(rect=[0, 0, 1, 0.95])
    st.pyplot(fig)


# ============================================================
# SIDEBAR MENU
# ============================================================
st.sidebar.title("🏭 UCG Platform v7.0")
st.sidebar.markdown("---")

menu = st.sidebar.selectbox(
    "📂 Menu",
    ["Dashboard", "Simulation", "Monte Carlo UQ", "Info", "Help", "About",
     "Patent Report", "ISO Report"]
)

st.sidebar.markdown("---")
st.sidebar.subheader("⚙️ Parametrlar")

coal_name = st.sidebar.selectbox("Ko'mir turi", list(COAL_DATABASE.keys()), index=0)
n_steps = st.sidebar.slider("Simulyatsiya qadamlari", 50, 500, 200, 10)
dt = st.sidebar.slider("Vaqt qadami (dt)", 0.1, 5.0, 1.0, 0.1)
T0 = st.sidebar.slider("Boshlang'ich harorat (K)", 600, 1800, 1200, 50)
P0 = st.sidebar.slider("Boshlang'ich bosim (MPa)", 0.5, 30.0, 10.0, 0.5)

st.sidebar.markdown("---")
st.sidebar.markdown(f"**Versiya:** 7.0.0")
st.sidebar.markdown(f"**Sana:** {datetime.now().strftime('%Y-%m-%d %H:%M')}")
st.sidebar.markdown("**Status:** Patent Pending")


# ============================================================
# MAIN APP PAGES
# ============================================================

if menu == "Dashboard":
    st.title("📊 UCG Digital Twin — Dashboard")
    st.markdown("---")

    # Simulyatsiyani ishga tushirish
    with st.spinner("Simulyatsiya hisoblanmoqda..."):
        df, engine, coal = run_simulation(coal_name, n_steps, dt, T0, P0)

    # KPI metrikalari
    final = df.iloc[-1]
    k1, k2, k3, k4, k5, k6 = st.columns(6)
    k1.metric("Konversiya (%)", f"{final['char_conversion']*100:.2f}")
    k2.metric("Harorat (K)", f"{final['temperature']:.1f}")
    k3.metric("Porozlik", f"{final['porosity']:.4f}")
    k4.metric("Novelty", f"{final.get('novelty_index', 0):.1f}%")
    k5.metric("Bosim (MPa)", f"{final['pressure']:.2f}")
    k6.metric("O'tkaz. (m²)", f"{final['permeability']:.2e}")

    st.markdown("---")

    # 2x3 asosiy grafiklar
    st.subheader("📈 Asosiy Grafiklar")
    col1, col2 = st.columns(2)
    with col1:
        plot_coal_conversion(df)
        plot_gas_composition(df)
        plot_temperature(df)
    with col2:
        plot_heat_source(df)
        plot_porosity(df)
        plot_pressure(df)

    st.markdown("---")

    # Qo'shimcha grafiklar
    st.subheader("🔬 Qo'shimcha Tahlil")
    col3, col4 = st.columns(2)
    with col3:
        plot_reaction_rates(df)
        plot_delta_H_bar()
        plot_ucg_zones(df)
    with col4:
        plot_permeability(df)
        plot_heat_share_pie(df)
        plot_novelty_index(df)

    st.markdown("---")

    # 3D Surface
    st.subheader("🌐 3D Surface Grafik")
    plot_3d_surface(df)


elif menu == "Simulation":
    st.title("🔬 Simulyatsiya Parametrlari")

    st.subheader("Ko'mir turi ma'lumotlari")
    coal = COAL_DATABASE[coal_name]
    info_col1, info_col2 = st.columns(2)
    with info_col1:
        st.write(f"**Nomi:** {coal.name}")
        st.write(f"**C:** {coal.carbon_frac*100:.1f}%  |  **H:** {coal.hydrogen_frac*100:.1f}%  |  **O:** {coal.oxygen_frac*100:.1f}%")
        st.write(f"**N:** {coal.nitrogen_frac*100:.2f}%  |  **S:** {coal.sulfur_frac*100:.2f}%")
    with info_col2:
        st.write(f"**Namlik:** {coal.moisture*100:.1f}%  |  **Kul:** {coal.ash*100:.1f}%")
        st.write(f"**Uchuvchi:** {coal.volatile_matter*100:.1f}%  |  **Qattiq C:** {coal.fixed_carbon*100:.1f}%")
        st.write(f"**HHV:** {coal.hhv_mj_kg:.1f} MJ/kg  |  **Reaktivlik:** {coal.reactivity_coeff}")

    st.markdown("---")
    st.subheader("Reaksiya Kinetikasi")
    rxn_data = []
    for name, rxn in UCGConfig.REACTIONS.items():
        rxn_data.append({"Reaksiya": name, "A": f"{rxn.A:.1e}", "Ea (kJ/mol)": f"{rxn.Ea/1000:.0f}",
                          "ΔH (kJ/mol)": f"{rxn.dH/1000:.1f}", "ΔS (J/mol·K)": f"{rxn.dS:.1f}",
                          "ΔG_ref (kJ/mol)": f"{rxn.dG_ref/1000:.1f}"})
    st.dataframe(pd.DataFrame(rxn_data), use_container_width=True)

    st.markdown("---")
    st.subheader("Simulyatsiyani Ishga Tushirish")
    if st.button("▶️ Simulyatsiyani Boshlash", type="primary"):
        with st.spinner("Hisoblanmoqda..."):
            df, engine, coal = run_simulation(coal_name, n_steps, dt, T0, P0)

        st.success(f"Simulyatsiya yakunlandi! {n_steps} qadam, dt={dt}")

        tab1, tab2, tab3 = st.tabs(["Konversiya", "Gaz Tarkibi", "Harorat"])
        with tab1:
            plot_coal_conversion(df)
        with tab2:
            plot_gas_composition(df)
        with tab3:
            plot_temperature(df)

        st.subheader("📋 Natijalar Jadvali")
        st.dataframe(df.tail(10), use_container_width=True)


elif menu == "Monte Carlo UQ":
    st.title("🎲 Monte Carlo Noaniqlikni Baholash")

    st.markdown("Simulyatsiya parametrlariga noaniqlik kiritib, ishonchli oraliqlarni (95% CI) hisoblash.")

    mc_col1, mc_col2, mc_col3 = st.columns(3)
    with mc_col1:
        n_sim = st.number_input("Simulyatsiya soni", min_value=100, max_value=5000, value=1000, step=100)
    with mc_col2:
        mc_steps = st.number_input("Qadamlar soni", min_value=20, max_value=200, value=100, step=10)
    with mc_col3:
        uncertainty = st.slider("Noaniqlik (σ)", 0.01, 0.30, 0.10, 0.01)

    if st.button("▶️ Monte Carlo ni Ishga Tushirish", type="primary"):
        with st.spinner(f"{n_sim} ta simulyatsiya ishga tushmoqda... Bu biroz vaqt olishi mumkin."):
            mc_summary, n_sim_actual = run_monte_carlo(T0, P0, coal_name, n_sim, mc_steps, dt, uncertainty)

        st.success(f"Monte Carlo yakunlandi! {n_sim_actual} ta simulyatsiya")

        plot_monte_carlo(mc_summary, n_sim_actual)

        # Natijalar jadvali
        st.subheader("📊 Yakuniy Noaniqlik Natijalari")
        final_idx = mc_steps - 1
        mc_table = []
        for key in ['temperature', 'char_conversion', 'porosity', 'novelty_index']:
            s = mc_summary[key]
            mc_table.append({
                "Parametr": key,
                "O'rtacha": f"{s['mean'][final_idx]:.4f}",
                "Std": f"{s['std'][final_idx]:.4f}",
                "95% CI Low": f"{s['ci95_low'][final_idx]:.4f}",
                "95% CI High": f"{s['ci95_high'][final_idx]:.4f}",
            })
        st.dataframe(pd.DataFrame(mc_table), use_container_width=True)


elif menu == "Info":
    st.title("ℹ️ UCG Platform haqida ma'lumot")

    st.header("UCG — Underground Coal Gasification")
    st.markdown("""
    **Yer osti ko'mir gazifikatsiyasi (UCG)** — bu ko'mirni yer ostida gazga aylantirish texnologiyasi.
    Ko'mir qatlami ichida maxsus chuqurlar orqali havo/kislorod/bug' yuborilib, ko'mir
    kimyoviy reaksiyalar natijasida gazga (CO, H₂, CH₄) aylanadi.

    ### Asosiy reaksiyalar:
    - **Oksidlash:** C + O₂ → CO₂ (ekzotermik, ΔH = -393.5 kJ/mol)
    - **Buduar:** C + CO₂ → 2CO (endotermik, ΔH = +172.0 kJ/mol)
    - **Bug' gazifikatsiyasi:** C + H₂O → CO + H₂ (endotermik, ΔH = +131.0 kJ/mol)
    - **Metanlash:** CO + 3H₂ → CH₄ + H₂O (ekzotermik, ΔH = -75.0 kJ/mol)
    - **WGS:** CO + H₂O → CO₂ + H₂ (ekzotermik, ΔH = -41.2 kJ/mol)
    - **Devolatilizatsiya:** C → H₂ + CH₄ + CO (endotermik)

    ### Platform imkoniyatlari:
    - 6 ta kimyoviy reaksiya moduli (Arrhenius kinetikasi)
    - Gibbs Free Energy hisoblash
    - Species Mass Balance (massa saqlanish qonuni)
    - Harorat va bosim evolyutsiyasi
    - Monte Carlo noaniqlik tahlili
    - KPI Dashboard
    - Patent hisobotlari eksport
    """)

    st.header("Ilmiy Adabiyotlar")
    refs = [
        "Biot, M.A. (1941). General theory of 3D consolidation. J. Appl. Phys., 12(2), 155-164.",
        "Hoek, E., & Brown, E.T. (2018). The Hoek-Brown failure criterion – 2018 edition. JRMGE.",
        "Yang, D. (2010). Stability of UCG. PhD Thesis, TU Delft.",
        "Perkins, G. (2018). Underground coal gasification – Part I. Progress in Energy and Combustion Science.",
        "JCGM 100:2008 (GUM). Evaluation of measurement data.",
    ]
    for r in refs:
        st.markdown(f"- {r}")


elif menu == "Help":
    st.title("❓ Yordam")

    st.markdown("""
    ### Qanday ishlatish?

    1. **Chap paneldan** ko'mir turini tanlang
    2. **Simulyatsiya parametrlarini** sozlang (qadamlar, dt, harorat, bosim)
    3. **Dashboard** sahifasida natijalarni ko'ring
    4. **Monte Carlo** sahifasida noaniqlik tahlilini o'tkazing
    5. **Patent Report** sahifasida hisobotlarni yuklab oling

    ### Grafiklar haqida

    | Rang | Ma'no |
    |------|-------|
    | 🔴 Qizil | Yuqori harorat / Ekzotermik / CO |
    | 🔵 Ko'k | Past harorat / Endotermik / H₂ |
    | 🟢 Yashil | Xavfsiz zona / CO₂ |
    | 🟡 Sariq | O'rta zona / Diqqat |
    | 🟣 Binafsha | CH₄ / Bosim |

    ### Zona ranglari

    - 🔴 **Past** — Boshlang'ich bosqich, kam konversiya
    - 🟡 **O'rta** — Faol gazifikatsiya, diqqat talab qilinadi
    - 🟢 **Yuqori** — To'liq gazifikatsiya, barqaror holat
    """)


elif menu == "About":
    st.title("📌 About")

    st.markdown("""
    ## UCG SCI-Grade Platform v7.0.0

    **Muallif:** Saitov Dilshodbek
    **Status:** Patent Pending (UzPatent + WIPO PCT)
    **Litsenziya:** Patent Pending — O'zbekiston 00XXXX + WIPO

    ### Texnologiyalar:
    - Python 3.12
    - Streamlit
    - NumPy / SciPy / Matplotlib
    - Arrhenius Kinetikasi
    - Monte Carlo UQ
    - Gibbs Free Energy
    - Species Mass Balance

    ### Versiya tarixi:
    | Versiya | Sana | O'zgarishlar |
    |---------|------|-------------|
    | v4.0.1 | 2026-06-21 | Patent-ready core |
    | v5.0.0 | 2026-06-21 | 20 critical fixes |
    | v6.1.0 | 2026-06-22 | 20 patent-grade fixes |
    | v7.0.0 | 2026-06-23 | Streamlit + 12 grafik + Monte Carlo + Gibbs |
    """)


elif menu == "Patent Report":
    st.title("📄 Patent Hisoboti")

    with st.spinner("Simulyatsiya hisoblanmoqda..."):
        df, engine, coal = run_simulation(coal_name, n_steps, dt, T0, P0)

    final = df.iloc[-1]

    st.subheader("Patentability Natijalari")
    p1, p2, p3, p4 = st.columns(4)
    p1.metric("Novelty Index", f"{final.get('novelty_index', 0):.1f}%")
    p2.metric("Konversiya", f"{final['char_conversion']*100:.2f}%")
    p3.metric("Porozlik", f"{final['porosity']:.4f}")
    p4.metric("Harorat (K)", f"{final['temperature']:.1f}")

    st.markdown("---")

    st.subheader("Gibbs Free Energy (Yakuniy)")
    gibbs_data = GibbsEnergyCalculator.compute_all(UCGConfig(), final['temperature'])
    gibbs_df = pd.DataFrame([
        {"Reaksiya": k, "ΔG (kJ/mol)": f"{v/1000:.1f}", "Spontan": "✅ Ha" if v < 0 else "❌ Yo'q"}
        for k, v in gibbs_data.items()
    ])
    st.dataframe(gibbs_df, use_container_width=True)

    st.markdown("---")

    st.subheader("Grafiklar")
    tab_a, tab_b, tab_c = st.tabs(["Konversiya + Issiqlik", "Gaz + Porozlik", "3D Surface"])
    with tab_a:
        plot_coal_conversion(df)
        plot_heat_source(df)
    with tab_b:
        plot_gas_composition(df)
        plot_porosity(df)
    with tab_c:
        plot_3d_surface(df)

    st.markdown("---")

    # Eksport
    st.subheader("📥 Eksport")
    export_col1, export_col2 = st.columns(2)

    if export_col1.button("📄 DOCX Hisobotni Yuklash"):
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            doc = Document()
            doc.add_heading('UCG Platform v7.0 — Patent Hisoboti', level=0)
            doc.add_paragraph(f'Sana: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f"Ko'mir turi: {coal.name}")
            doc.add_heading('KPI Natijalari', level=1)
            for kpi in ['char_conversion', 'temperature', 'porosity', 'novelty_index']:
                if kpi in df.columns:
                    doc.add_paragraph(f'{kpi}: {final[kpi]}', style='List Bullet')
            doc.add_heading('Gibbs Free Energy', level=1)
            for k, v in gibbs_data.items():
                doc.add_paragraph(f'{k}: ΔG = {v/1000:.1f} kJ/mol', style='List Bullet')
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button("⬇️ DOCX yuklab olish", buf.getvalue(),
                               file_name="ucg_v7_patent_report.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except ImportError:
            st.error("python-docx o'rnatilmagan!")

    if export_col2.button("📋 JSON Hisobotni Yuklash"):
        report = {"version": "7.0.0", "timestamp": datetime.now().isoformat(),
                  "coal_type": coal.name, "final_state": {k: float(v) for k, v in final.items() if not isinstance(v, dict)},
                  "gibbs": {k: v for k, v in gibbs_data.items()}}
        json_str = json.dumps(report, ensure_ascii=False, indent=2, default=str)
        st.download_button("⬇️ JSON yuklab olish", json_str,
                           file_name="ucg_v7_report.json", mime="application/json")


elif menu == "ISO Report":
    st.title("📋 ISO/ISRM Muvofiqlik Hisoboti")

    st.subheader("ISO Muvofiqlik Matritsasi")
    iso_data = pd.DataFrame([
        {"Standart": "ISO 9001", "Soxa": "Sifat boshqaruvi", "Status": "Mapped",
         "Dalil": "Versioned report workflow, verification, audit trail"},
        {"Standart": "ISO 31000", "Soxa": "Xavf boshqaruvi", "Status": "Mapped",
         "Dalil": "Monte Carlo UQ, risk index, sensitivity analysis"},
        {"Standart": "ISO 27001", "Soxa": "Axborot xavfsizligi", "Status": "Mapped",
         "Dalil": "SHA-256 hash chain, RSA-4096 signature, WORM storage"},
        {"Standart": "IEC 61508", "Soxa": "Funksional xavfsizlik", "Status": "Partial",
         "Dalil": "Alarm logic, monitoring architecture"},
        {"Standart": "ISRM", "Soxa": "Tosh mexanikasi", "Status": "Mapped",
         "Dalil": "Hoek-Brown, UCS/GSI, verification workflow"},
    ])
    st.dataframe(iso_data, use_container_width=True)

    st.markdown("---")

    st.subheader("ISO 9001 — Sifat Boshqaruvi Tahlili")
    st.markdown("""
    - ✅ **Hujjat boshqaruvi:** Versiyalangan hisobotlar, SHA-256 traceability
    - ✅ **Sifat siyosati:** Ilmiy tasdiqlash bosqichlari (5-stage validation)
    - ✅ **Xavf asosida fikrlash:** Monte Carlo noaniqlik tahlili
    - ⚠️ **Ichki audit:** Avtomatik audit trail mavjud, lekin mustaqil audit rejalashtirish yo'q
    """)

    st.subheader("ISO 31000 — Xavf Boshqaruvi Tahlili")
    st.markdown("""
    - ✅ **Xavf identifikatsiyasi:** FOS < 1.0 aniqlash, Monte Carlo PF hisoblash
    - ✅ **Xavf baholash:** 95% CI, sensitivity analysis (Sobol/Morris)
    - ✅ **Xavf davolash:** Tavsiyalar avtomatik generatsiya
    - ⚠️ **Xavf ishtahashi:** Xavf ishtahashi bayonoti mavjud emas
    """)

    st.subheader("ISO 27001 — Axborot Xavfsizligi Tahlili")
    st.markdown("""
    - ✅ **Axborot xavfsizligi siyosati:** SHA-256 + RSA-4096 imzolash
    - ✅ **Kirish nazorati:** .env fayl, credential vault (HashiCorp/Azure Key Vault)
    - ✅ **Inssident boshqaruvi:** Blockchain audit trail (immutable hash chain)
    - ⚠️ **Inssident javob rejasini:** Hujjatlashtirilmagan
    """)
